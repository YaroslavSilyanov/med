import os
import docx
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

class DocumentGenerator:
    def __init__(self):
        """Инициализация генератора документов"""
        self.documents_path = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))), 'resources', 'documents')
        os.makedirs(self.documents_path, exist_ok=True)
    
    def generate_analysis_report(self, analysis_data):
        """
        Генерировать отчет по анализу в формате Word
        
        Args:
            analysis_data (dict): Данные анализа
        
        Returns:
            str: Путь к сгенерированному файлу
        """
        # Создать новый документ
        doc = docx.Document()
        
        # Установить поля страницы
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(3)
            section.right_margin = Cm(2)
        
        # Добавить заголовок
        header = doc.add_heading(f'РЕЗУЛЬТАТЫ АНАЛИЗА', level=1)
        header_format = header.runs[0].font
        header_format.size = Pt(16)
        header_format.bold = True
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Информация о пациенте
        patient_info = doc.add_paragraph()
        patient_info.alignment = WD_ALIGN_PARAGRAPH.LEFT
        patient_run = patient_info.add_run(f"Пациент: {analysis_data['patient']['full_name']}")
        patient_run.font.size = Pt(12)
        patient_run.font.bold = True
        
        birth_date_run = patient_info.add_run(f"\nДата рождения: {analysis_data['patient']['birth_date']}")
        birth_date_run.font.size = Pt(12)
        
        # Информация об анализе
        analysis_info = doc.add_paragraph()
        analysis_info.alignment = WD_ALIGN_PARAGRAPH.LEFT
        analysis_type_run = analysis_info.add_run(f"Тип анализа: {analysis_data['analysis_type']['name']}")
        analysis_type_run.font.size = Pt(12)
        analysis_type_run.font.bold = True
        
        date_run = analysis_info.add_run(f"\nДата взятия: {analysis_data['date_taken']}")
        date_run.font.size = Pt(12)
        
        lab_tech_run = analysis_info.add_run(f"\nЛаборант: {analysis_data['lab_technician']}")
        lab_tech_run.font.size = Pt(12)
        
        # Таблица с результатами
        doc.add_paragraph()
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # Заголовки таблицы
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Параметр'
        hdr_cells[1].text = 'Результат'
        hdr_cells[2].text = 'Единицы измерения'
        hdr_cells[3].text = 'Норма'
        hdr_cells[4].text = 'Оценка'
        
        # Сделать заголовки жирными
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(11)
        
        # Добавить данные в таблицу
        for param in analysis_data['parameters']:
            row_cells = table.add_row().cells
            row_cells[0].text = param['name']
            row_cells[1].text = str(param['value'])
            row_cells[2].text = param['unit'] if param['unit'] else ''
            
            # Отображение нормы
            if param['normal_min'] is not None and param['normal_max'] is not None:
                row_cells[3].text = f"{param['normal_min']} - {param['normal_max']}"
            else:
                row_cells[3].text = "Не определена"
            
            # Оценка результата
            if param['is_normal'] is not None:
                if param['is_normal']:
                    row_cells[4].text = "В норме"
                    # Установим зеленый цвет для нормы
                    for paragraph in row_cells[4].paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(0, 128, 0)  # Зеленый
                else:
                    row_cells[4].text = "Отклонение"
                    # Установим красный цвет для отклонения
                    for paragraph in row_cells[4].paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(255, 0, 0)  # Красный
            else:
                row_cells[4].text = "Не определено"
        
        # Добавить заключение
        doc.add_paragraph()
        conclusion = doc.add_paragraph()
        conclusion.alignment = WD_ALIGN_PARAGRAPH.LEFT
        conclusion_run = conclusion.add_run("Заключение:")
        conclusion_run.font.bold = True
        conclusion_run.font.size = Pt(12)
        
        # Автоматическое заключение на основе результатов
        abnormal_params = [p for p in analysis_data['parameters'] if p['is_normal'] is not None and not p['is_normal']]
        if abnormal_params:
            conclusion_text = f"\nОбнаружены отклонения в следующих показателях: {', '.join([p['name'] for p in abnormal_params])}."
            conclusion_text += "\nРекомендуется консультация врача."
        else:
            conclusion_text = "\nВсе показатели в пределах нормы."
        
        conclusion.add_run(conclusion_text).font.size = Pt(12)
        
        # Добавить подпись
        doc.add_paragraph()
        doc.add_paragraph("Подпись врача: _______________________").alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Имя файла
        filename = f"Анализ_{analysis_data['analysis_type']['name']}_{analysis_data['patient']['full_name']}_{analysis_data['date_taken']}.docx"
        # Заменим недопустимые символы в имени файла
        filename = filename.replace(':', '_').replace('/', '_').replace('\\', '_')
        
        # Путь к файлу
        file_path = os.path.join(self.documents_path, filename)
        
        # Сохранить документ
        doc.save(file_path)
        
        return file_path

    def generate_patient_card(self, patient_data, analysis_results):
        """
        Генерировать карточку пациента в формате Word
        
        Args:
            patient_data (dict): Данные о пациенте
            analysis_results (list): Список анализов пациента
        
        Returns:
            str: Путь к сгенерированному файлу
        """
        # Создать новый документ
        doc = docx.Document()
        
        # Установить поля страницы
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(3)
            section.right_margin = Cm(2)
        
        # Добавить заголовок
        header = doc.add_heading(f'КАРТА ПАЦИЕНТА', level=1)
        header_format = header.runs[0].font
        header_format.size = Pt(16)
        header_format.bold = True
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Информация о пациенте
        patient_info = doc.add_paragraph()
        patient_info.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        patient_run = patient_info.add_run(f"ФИО: {patient_data['full_name']}")
        patient_run.font.size = Pt(12)
        patient_run.font.bold = True
        
        patient_info.add_run(f"\nДата рождения: {patient_data['birth_date']}").font.size = Pt(12)
        patient_info.add_run(f"\nПол: {patient_data['gender']}").font.size = Pt(12)
        
        if patient_data['phone_number']:
            patient_info.add_run(f"\nТелефон: {patient_data['phone_number']}").font.size = Pt(12)
        
        if patient_data['email']:
            patient_info.add_run(f"\nEmail: {patient_data['email']}").font.size = Pt(12)
        
        if patient_data['address']:
            patient_info.add_run(f"\nАдрес: {patient_data['address']}").font.size = Pt(12)
        
        # Раздел с анализами
        if analysis_results:
            doc.add_paragraph()
            analysis_header = doc.add_heading('История анализов', level=2)
            analysis_header.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Таблица с анализами
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            
            # Заголовки таблицы
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Дата'
            hdr_cells[1].text = 'Тип анализа'
            hdr_cells[2].text = 'Статус'
            
            # Сделать заголовки жирными
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(11)
            
            # Добавить данные в таблицу
            for result in analysis_results:
                row_cells = table.add_row().cells
                row_cells[0].text = result['date_taken']
                row_cells[1].text = result['analysis_type']
                row_cells[2].text = result['status']
        else:
            doc.add_paragraph("История анализов отсутствует").font.italic = True
        
        # Имя файла
        filename = f"Карта_{patient_data['full_name']}.docx"
        # Заменим недопустимые символы в имени файла
        filename = filename.replace(':', '_').replace('/', '_').replace('\\', '_')
        
        # Путь к файлу
        file_path = os.path.join(self.documents_path, filename)
        
        # Сохранить документ
        doc.save(file_path)
        
        return file_path
    
    def generate_appointment_referral(self, patient_data, doctor_data, appointment_data, analysis_results=None):
        """
        Генерировать направление на прием к врачу в формате Word
        
        Args:
            patient_data (dict): Данные о пациенте
            doctor_data (dict): Данные о враче
            appointment_data (dict): Данные о приеме
            analysis_results (list, optional): Список анализов пациента
        
        Returns:
            str: Путь к сгенерированному файлу
        """
        # Создать новый документ
        doc = docx.Document()
        
        # Установить поля страницы
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(3)
            section.right_margin = Cm(2)
        
        # Добавить заголовок
        header = doc.add_heading(f'НАПРАВЛЕНИЕ НА ПРИЕМ', level=1)
        header_format = header.runs[0].font
        header_format.size = Pt(16)
        header_format.bold = True
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Информация о дате
        date_info = doc.add_paragraph()
        date_info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        current_date = datetime.now().strftime("%d.%m.%Y")
        date_run = date_info.add_run(f"Дата выдачи: {current_date}")
        date_run.font.size = Pt(10)
        
        # Информация о пациенте
        doc.add_paragraph()
        patient_title = doc.add_heading('Пациент:', level=2)
        patient_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        patient_info = doc.add_paragraph()
        patient_info.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        patient_run = patient_info.add_run(f"ФИО: {patient_data['full_name']}")
        patient_run.font.size = Pt(12)
        patient_run.font.bold = True
        
        patient_info.add_run(f"\nДата рождения: {patient_data['birth_date']}").font.size = Pt(12)
        patient_info.add_run(f"\nПол: {patient_data['gender']}").font.size = Pt(12)
        
        if patient_data.get('phone_number'):
            patient_info.add_run(f"\nТелефон: {patient_data['phone_number']}").font.size = Pt(12)
        
        # Информация о враче
        doc.add_paragraph()
        doctor_title = doc.add_heading('Врач:', level=2)
        doctor_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        doctor_info = doc.add_paragraph()
        doctor_info.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        doctor_run = doctor_info.add_run(f"ФИО: {doctor_data['full_name']}")
        doctor_run.font.size = Pt(12)
        doctor_run.font.bold = True
        
        doctor_info.add_run(f"\nСпециализация: {doctor_data.get('specialization', 'Не указана')}").font.size = Pt(12)
        
        # Информация о приеме
        doc.add_paragraph()
        appointment_title = doc.add_heading('Информация о приеме:', level=2)
        appointment_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        appointment_info = doc.add_paragraph()
        appointment_info.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Разделим дату и время, если они в одной строке
        appointment_datetime = appointment_data.get('appointment_date', '')
        if ' ' in appointment_datetime:
            date_part, time_part = appointment_datetime.split(' ', 1)
        else:
            date_part = appointment_datetime
            time_part = ''
        
        appointment_info.add_run(f"Дата приема: {date_part}").font.size = Pt(12)
        appointment_info.add_run(f"\nВремя приема: {time_part}").font.size = Pt(12)
        
        status = appointment_data.get('status', 'scheduled')
        status_text = {
            'scheduled': 'Запланирован',
            'completed': 'Завершен',
            'cancelled': 'Отменен'
        }.get(status, status)
        
        appointment_info.add_run(f"\nСтатус: {status_text}").font.size = Pt(12)
        
        if appointment_data.get('notes'):
            notes_run = appointment_info.add_run(f"\nПримечания: {appointment_data['notes']}")
            notes_run.font.size = Pt(12)
            notes_run.italic = True
        
        # Если есть результаты анализов, добавим их
        if analysis_results:
            doc.add_paragraph()
            analysis_title = doc.add_heading('Результаты анализов:', level=2)
            analysis_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Таблица с анализами
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            
            # Заголовки таблицы
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Дата'
            hdr_cells[1].text = 'Тип анализа'
            hdr_cells[2].text = 'Статус'
            
            # Сделать заголовки жирными
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(11)
            
            # Добавить данные в таблицу
            for result in analysis_results:
                row_cells = table.add_row().cells
                row_cells[0].text = result['date_taken']
                row_cells[1].text = result['analysis_type']
                row_cells[2].text = result['status']
        
        # Добавить подпись
        doc.add_paragraph()
        signature = doc.add_paragraph()
        signature.alignment = WD_ALIGN_PARAGRAPH.LEFT
        signature.add_run("Подпись врача: _______________________").font.size = Pt(12)
        
        # Добавить печать
        signature.add_run("\n\nМ.П.").font.size = Pt(12)
        
        # Имя файла
        appointment_date_formatted = date_part.replace('.', '_').replace('/', '_').replace('\\', '_')
        filename = f"Направление_{patient_data['full_name']}_{appointment_date_formatted}.docx"
        # Заменим недопустимые символы в имени файла
        filename = filename.replace(':', '_').replace('/', '_').replace('\\', '_')
        
        # Путь к файлу
        file_path = os.path.join(self.documents_path, filename)
        
        # Сохранить документ
        doc.save(file_path)
        
        return file_path 