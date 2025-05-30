from PySide6.QtWidgets import (QMainWindow, QWidget, QLabel, QComboBox, QPushButton,
                               QVBoxLayout, QHBoxLayout, QMessageBox, QFormLayout, 
                               QTableWidget, QTableWidgetItem, QLineEdit, QDialog,
                               QTabWidget, QGroupBox, QScrollArea, QFrame, QListWidget,
                               QListWidgetItem, QGridLayout, QDateEdit, QSpinBox,
                               QRadioButton, QButtonGroup, QCheckBox, QTextEdit,
                               QHeaderView, QStackedWidget, QSplitter, QTimeEdit,
                               QFileDialog)
from PySide6.QtCore import Qt, Signal, QDate, QSize, QTimer, QTime
from PySide6.QtGui import QFont, QIcon, QColor, QPixmap, QPainter, QPen, QBrush, QPainterPath
from datetime import datetime, timedelta
import sys
import json
import os
import tempfile
import docx
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
import subprocess
import platform
import xlwt
import ssl
import report_generator

from database_connection import db

# Попытка импорта классов из app.utils
try:
    from app.utils.document_generator import DocumentGenerator
    from app.utils.email_sender import EmailSender
    print("Импортированы классы из app.utils")
except ImportError:
    print("Невозможно импортировать классы из app.utils, используются встроенные версии")
    # DocumentGenerator класс определен в другом месте - app.utils.document_generator


class EmailSender:
    """Класс для отправки электронных писем"""
    
    def __init__(self, smtp_server='smtp.gmail.com', port=587, test_mode=False):
        """Инициализация параметров подключения к SMTP серверу"""
        self.smtp_server = smtp_server
        self.port = port
        # Эти данные должны быть настроены в реальном приложении
        self.username = "medcenter.example@gmail.com"  # Заменить на реальный адрес
        self.app_password = "app_password_here"         # Заменить на пароль приложения Gmail
        self.test_mode = test_mode
    
    def send_email(self, recipient_email, subject, message, attachments=None):
        """
        Отправка электронного письма
        
        :param recipient_email: Email получателя
        :param subject: Тема письма
        :param message: Текст письма (может быть в формате HTML)
        :param attachments: Список путей к файлам для прикрепления
        :return: True в случае успеха, False в случае ошибки
        """
        try:
            # Создание объекта сообщения
            from email.mime.multipart import MIMEMultipart
            from email.mime.text import MIMEText
            from email.mime.application import MIMEApplication
            import os
            
            msg = MIMEMultipart()
            msg['From'] = self.username
            msg['To'] = recipient_email
            msg['Subject'] = subject
            
            # Добавление текста письма
            msg.attach(MIMEText(message, 'html'))
            
            # Добавление вложений
            if attachments:
                for file_path in attachments:
                    if os.path.isfile(file_path):
                        with open(file_path, 'rb') as file:
                            part = MIMEApplication(file.read(), Name=os.path.basename(file_path))
                            part['Content-Disposition'] = f'attachment; filename="{os.path.basename(file_path)}"'
                            msg.attach(part)

            # В тестовом режиме не отправляем письмо, а только выводим информацию
            if self.test_mode:
                print(f"\n[ТЕСТОВЫЙ РЕЖИМ] Отправка email:")
                print(f"От: {self.username}")
                print(f"Кому: {recipient_email}")
                print(f"Тема: {subject}")
                print(f"Содержание: HTML-письмо")
                if attachments:
                    print(f"Вложения: {', '.join([os.path.basename(f) for f in attachments if os.path.isfile(f)])}")
                print("Email успешно отправлен (тестовый режим)")
                return True
            
            # Отправка письма
            import smtplib
            
            context = ssl.create_default_context()
            with smtplib.SMTP(self.smtp_server, self.port) as server:
                server.starttls(context=context)
                server.login(self.username, self.app_password)
                server.sendmail(self.username, recipient_email, msg.as_string())
            
            return True
            
        except Exception as e:
            print(f"Ошибка при отправке email: {str(e)}")
            return False
    
    def send_report(self, recipient_email, recipient_name, report_type, report_period, report_file_path, additional_text=None):
        """
        Отправка отчета по электронной почте сотруднику
        
        :param recipient_email: Email получателя (сотрудника)
        :param recipient_name: Имя получателя (сотрудника)
        :param report_type: Тип отчета
        :param report_period: Период отчета
        :param report_file_path: Путь к файлу отчета для прикрепления
        :param additional_text: Дополнительный текст для включения в письмо
        :return: True в случае успеха, False в случае ошибки
        """
        try:
            # Проверка существования файла отчета
            if not os.path.isfile(report_file_path):
                print(f"Ошибка: файл отчета {report_file_path} не найден")
                return False
            
            # Создание HTML-содержимого письма
            html_content = f"""
            <html>
                <head>
                    <style>
                        body {{ font-family: Arial, sans-serif; line-height: 1.6; }}
                        .container {{ width: 80%; margin: 0 auto; padding: 20px; }}
                        h1 {{ color: #2c3e50; }}
                        .report-info {{ background-color: #f8f9fa; padding: 15px; border-radius: 5px; margin: 20px 0; }}
                        .additional-info {{ background-color: #e8f4fe; padding: 10px; border-left: 4px solid #2196f3; margin-top: 20px; }}
                    </style>
                </head>
                <body>
                    <div class="container">
                        <h1>Отчет: {report_type}</h1>
                        <p>Уважаемый(ая) <strong>{recipient_name}</strong>,</p>
                        <p>Направляем Вам отчет <strong>"{report_type}"</strong> за период <strong>{report_period}</strong>.</p>
                        
                        <div class="report-info">
                            <p>Файл отчета прикреплен к письму.</p>
                            <p>Тип файла: {os.path.splitext(report_file_path)[1]}</p>
                        </div>
            """
            
            # Добавление дополнительной информации, если она указана
            if additional_text:
                html_content += f"""
                        <div class="additional-info">
                            <p><strong>Дополнительная информация:</strong></p>
                            <p>{additional_text}</p>
                        </div>
                """
            
            html_content += """
                        <p>С уважением,<br>Медицинский центр</p>
                    </div>
                </body>
            </html>
            """
            
            # Отправка email с отчетом
            return self.send_email(
                recipient_email=recipient_email,
                subject=f"Отчет: {report_type} за {report_period}",
                message=html_content,
                attachments=[report_file_path]
            )
            
        except Exception as e:
            print(f"Ошибка при отправке отчета: {str(e)}")
            return False


class DocumentGenerator:
    """Класс для создания документов в разных форматах"""
    
    def __init__(self):
        """Инициализация генератора документов"""
        pass
    
    def generate_analysis_report(self, result_details):
        """
        Генерация отчета с результатами анализов в формате Word
        
        :param result_details: Словарь с данными результатов анализа
        :return: Путь к созданному файлу
        """
        try:
            # Создание документа Word
            doc = docx.Document()
            
            # Настройка стилей
            style = doc.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(12)
            
            # Заголовок документа
            title = doc.add_heading('РЕЗУЛЬТАТЫ МЕДИЦИНСКОГО АНАЛИЗА', level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Информация о пациенте
            doc.add_heading('Информация о пациенте', level=2)
            
            patient_table = doc.add_table(rows=3, cols=2)
            patient_table.style = 'Table Grid'
            
            # Заполнение данных пациента
            patient = result_details['patient']
            
            cell = patient_table.cell(0, 0)
            cell.text = 'ФИО пациента:'
            
            cell = patient_table.cell(0, 1)
            cell.text = patient.get('full_name', '')
            
            cell = patient_table.cell(1, 0)
            cell.text = 'Дата рождения:'
            
            cell = patient_table.cell(1, 1)
            cell.text = patient.get('birth_date', '')
            
            cell = patient_table.cell(2, 0)
            cell.text = 'Полис ОМС:'
            
            cell = patient_table.cell(2, 1)
            cell.text = patient.get('insurance_policy', '')
            
            # Информация об анализе
            doc.add_paragraph('')  # Пустой параграф для отступа
            doc.add_heading('Информация об анализе', level=2)
            
            analysis_table = doc.add_table(rows=4, cols=2)
            analysis_table.style = 'Table Grid'
            
            # Заполнение данных анализа
            cell = analysis_table.cell(0, 0)
            cell.text = 'Тип анализа:'
            
            cell = analysis_table.cell(0, 1)
            cell.text = result_details['analysis_type'].get('name', '')
            
            cell = analysis_table.cell(1, 0)
            cell.text = 'Дата взятия пробы:'
            
            cell = analysis_table.cell(1, 1)
            cell.text = result_details.get('date_taken', '')
            
            cell = analysis_table.cell(2, 0)
            cell.text = 'Дата выполнения:'
            
            cell = analysis_table.cell(2, 1)
            cell.text = result_details.get('result_date', '')
            
            cell = analysis_table.cell(3, 0)
            cell.text = 'Лаборант:'
            
            cell = analysis_table.cell(3, 1)
            cell.text = result_details.get('lab_technician', '')
            
            # Результаты анализа
            doc.add_paragraph('')  # Пустой параграф для отступа
            doc.add_heading('Результаты', level=2)
            
            # Парсинг JSON результатов
            try:
                result_values = json.loads(result_details.get('results', '{}'))
            except json.JSONDecodeError:
                result_values = {}
            
            # Если результаты есть, создаем таблицу с ними
            if result_values:
                results_table = doc.add_table(rows=1, cols=3)
                results_table.style = 'Table Grid'
                
                # Заголовки таблицы
                header_cells = results_table.rows[0].cells
                header_cells[0].text = 'Параметр'
                header_cells[1].text = 'Значение'
                header_cells[2].text = 'Норма'
                
                # Заполнение результатов
                for param, value in result_values.items():
                    # Добавление новой строки
                    row_cells = results_table.add_row().cells
                    row_cells[0].text = param
                    row_cells[1].text = str(value)
                    
                    # Нормальные значения (для примера)
                    normal_values = self._get_normal_values(param)
                    row_cells[2].text = normal_values
            else:
                doc.add_paragraph('Нет данных о результатах анализа.')
            
            # Заключение
            doc.add_paragraph('')  # Пустой параграф для отступа
            doc.add_heading('Заключение', level=2)
            
            if result_details.get('conclusion'):
                doc.add_paragraph(result_details.get('conclusion'))
            else:
                doc.add_paragraph('Заключение отсутствует')
            
            # Дата и подпись
            doc.add_paragraph('')  # Пустой параграф для отступа
            doc.add_paragraph(f"Дата: {result_details.get('result_date', '')}")
            doc.add_paragraph('Подпись врача: ___________________')
            
            # Создание временного файла для сохранения документа
            file_name = f"analysis_result_{result_details.get('id')}.docx"
            file_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), file_name)
            
            # Сохранение документа
            doc.save(file_path)
            
            return file_path
            
        except Exception as e:
            print(f"Ошибка при создании документа: {str(e)}")
            return None
    
    def _get_normal_values(self, parameter):
        """
        Получение нормальных значений для параметра анализа
        
        :param parameter: Имя параметра анализа
        :return: Строка с нормальными значениями
        """
        # Словарь с нормальными значениями для некоторых параметров
        normal_values = {
            'Гемоглобин': '120-160 г/л',
            'Эритроциты': '3.8-5.5 млн/мкл',
            'Лейкоциты': '4.0-9.0 тыс/мкл',
            'Тромбоциты': '180-320 тыс/мкл',
            'СОЭ': '2-15 мм/ч',
            'Глюкоза': '3.9-6.1 ммоль/л',
            'Холестерин': '3.0-5.2 ммоль/л',
            'Билирубин': '3.4-17.1 мкмоль/л',
            'АЛТ': '5-40 ед/л',
            'АСТ': '5-40 ед/л',
            'Креатинин': '53-106 мкмоль/л',
            'Мочевина': '2.5-8.3 ммоль/л',
            'pH': '5.0-7.0',
            'Белок': 'Отсутствует',
            'Кетоновые тела': 'Отсутствуют'
        }
        
        return normal_values.get(parameter, 'Не указано')


class AddEditUserDialog(QDialog):
    """Диалоговое окно для добавления/редактирования пользователя"""
    
    def __init__(self, user_data=None, parent=None):
        super().__init__(parent)
        self.user_data = user_data  # None, если создаем нового пользователя
        
        self.setWindowTitle("Добавление пользователя" if user_data is None else "Редактирование пользователя")
        self.setMinimumWidth(400)
        self.setup_ui()
    
    def setup_ui(self):
        """Настройка интерфейса диалога"""
        layout = QVBoxLayout()
        
        # Форма с полями ввода
        form_layout = QFormLayout()
        
        # Имя пользователя
        self.username_input = QLineEdit()
        if self.user_data:
            self.username_input.setText(self.user_data.get('username', ''))
            # Запрещаем изменение логина только для администратора
            if self.user_data.get('username') == 'admin':
                self.username_input.setEnabled(False)  # Запрещаем изменение логина администратора
        form_layout.addRow("Логин:", self.username_input)
        
        # Пароль
        self.password_input = QLineEdit()
        self.password_input.setEchoMode(QLineEdit.Password)
        form_layout.addRow("Пароль:", self.password_input)
        
        # ФИО
        self.full_name_input = QLineEdit()
        if self.user_data:
            self.full_name_input.setText(self.user_data.get('full_name', ''))
        form_layout.addRow("ФИО:", self.full_name_input)
        
        # Email
        self.email_input = QLineEdit()
        if self.user_data:
            self.email_input.setText(self.user_data.get('email', ''))
        form_layout.addRow("Email:", self.email_input)
        
        # Роль
        self.role_combo = QComboBox()
        self.role_combo.addItem("Администратор", "admin")
        self.role_combo.addItem("Врач", "doctor")
        self.role_combo.addItem("Лаборант", "lab")
        
        if self.user_data:
            role = self.user_data.get('role', '')
            index = self.role_combo.findData(role)
            if index >= 0:
                self.role_combo.setCurrentIndex(index)
        
        form_layout.addRow("Роль:", self.role_combo)
        
        # Специализация (только для врачей)
        self.specialization_input = QLineEdit()
        self.specialization_input.setVisible(False)
        self.specialization_label = QLabel("Специализация:")
        self.specialization_label.setVisible(False)
        form_layout.addRow(self.specialization_label, self.specialization_input)
        
        # Статус (только для редактирования)
        if self.user_data:
            self.status_combo = QComboBox()
            self.status_combo.addItem("Активен", "active")
            self.status_combo.addItem("Заблокирован", "blocked")
            
            status = self.user_data.get('status', 'active')
            index = self.status_combo.findData(status)
            if index >= 0:
                self.status_combo.setCurrentIndex(index)
            
            form_layout.addRow("Статус:", self.status_combo)
        
        layout.addLayout(form_layout)
        
        # Кнопки
        buttons_layout = QHBoxLayout()
        
        save_button = QPushButton("Сохранить")
        save_button.clicked.connect(self.save_user)
        buttons_layout.addWidget(save_button)
        
        cancel_button = QPushButton("Отмена")
        cancel_button.clicked.connect(self.reject)
        buttons_layout.addWidget(cancel_button)
        
        layout.addLayout(buttons_layout)
        
        self.setLayout(layout)
        
        # Обработка изменения роли
        self.role_combo.currentIndexChanged.connect(self.on_role_changed)
        self.on_role_changed(self.role_combo.currentIndex())
    
    def on_role_changed(self, index):
        """Обработка изменения роли пользователя"""
        role = self.role_combo.itemData(index)
        
        # Показываем поле специализации только для врачей
        is_doctor = (role == "doctor")
        self.specialization_input.setVisible(is_doctor)
        self.specialization_label.setVisible(is_doctor)
        
        # Обновляем размер окна
        self.adjustSize()
    
    def save_user(self):
        """Сохранение данных пользователя"""
        # Валидация данных
        username = self.username_input.text().strip()
        password = self.password_input.text()
        full_name = self.full_name_input.text().strip()
        email = self.email_input.text().strip()
        role = self.role_combo.currentData()
        
        if not username:
            QMessageBox.warning(self, "Ошибка", "Логин не может быть пустым")
            return
        
        if not self.user_data and not password:
            QMessageBox.warning(self, "Ошибка", "Пароль не может быть пустым")
            return
        
        if not full_name:
            QMessageBox.warning(self, "Ошибка", "ФИО не может быть пустым")
            return
        
        # Проверка уникальности логина для нового пользователя
        if not self.user_data:
            user_check = db.fetch_one("SELECT id FROM users WHERE username = ?", (username,))
            if user_check:
                QMessageBox.warning(self, "Ошибка", "Пользователь с таким логином уже существует")
                return
        # Проверка уникальности логина при редактировании (если логин был изменен)
        elif self.user_data and username != self.user_data['username']:
            user_check = db.fetch_one("SELECT id FROM users WHERE username = ?", (username,))
            if user_check:
                QMessageBox.warning(self, "Ошибка", "Пользователь с таким логином уже существует")
                return
        
        try:
            if self.user_data:  # Редактирование существующего пользователя
                user_id = self.user_data['id']
                status = self.status_combo.currentData()
                
                # Защита от смены роли для главного администратора (admin)
                if self.user_data['username'] == 'admin' and role != 'admin':
                    QMessageBox.warning(self, "Ошибка", "Нельзя изменить роль главного администратора")
                    return
                
                # Базовый запрос на обновление
                if self.user_data['username'] == 'admin':
                    # Для администратора обновляем только основные данные без роли
                    query = """
                    UPDATE users 
                    SET full_name = ?, email = ?, status = ?
                    """
                    params = [full_name, email, status]
                else:
                    # Для всех остальных пользователей обновляем все данные включая роль
                    # Также обновляем логин, если он был изменен
                    if username != self.user_data['username']:
                        query = """
                        UPDATE users 
                        SET username = ?, full_name = ?, email = ?, role = ?, status = ?
                        """
                        params = [username, full_name, email, role, status]
                    else:
                        query = """
                        UPDATE users 
                        SET full_name = ?, email = ?, role = ?, status = ?
                        """
                        params = [full_name, email, role, status]
                
                # Добавляем пароль, если он был введен
                if password:
                    query += ", password = ?"
                    params.append(password)
                
                query += " WHERE id = ?"
                params.append(user_id)
                
                print(f"Выполнение обновления пользователя: {query}")
                print(f"Параметры: {params}")
                
                if db.execute_query(query, params):
                    # Обновление информации о враче, если это врач
                    if role == "doctor":
                        specialization = self.specialization_input.text().strip()
                        
                        # Проверяем, существует ли уже запись о враче
                        doctor = db.fetch_one("SELECT id FROM doctors WHERE user_id = ?", (user_id,))
                        
                        if doctor:
                            db.execute_query(
                                "UPDATE doctors SET specialization = ? WHERE user_id = ?",
                                (specialization, user_id)
                            )
                        else:
                            db.execute_query(
                                "INSERT INTO doctors (user_id, specialization) VALUES (?, ?)",
                                (user_id, specialization)
                            )
                    
                    QMessageBox.information(self, "Успех", "Пользователь успешно обновлен")
                    self.accept()
                else:
                    QMessageBox.warning(self, "Ошибка", "Не удалось обновить пользователя")
            else:  # Добавление нового пользователя
                # Добавление пользователя
                print(f"Добавление нового пользователя: {username}, {role}")
                user_id = db.add_user(username, password, full_name, role, email)
                
                if user_id:
                    # Добавление информации о враче, если это врач
                    if role == "doctor":
                        specialization = self.specialization_input.text().strip()
                        db.execute_query(
                            "INSERT INTO doctors (user_id, specialization) VALUES (?, ?)",
                            (user_id, specialization)
                        )
                    
                    QMessageBox.information(self, "Успех", "Пользователь успешно добавлен")
                    self.accept()
                else:
                    QMessageBox.warning(self, "Ошибка", "Не удалось добавить пользователя")
        
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Произошла ошибка: {str(e)}")


class UserListWidget(QWidget):
    """Виджет для отображения списка пользователей"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setup_ui()
    
    def setup_ui(self):
        """Настройка интерфейса"""
        layout = QVBoxLayout(self)
        
        # Верхняя панель с кнопками
        top_panel = QHBoxLayout()
        
        refresh_button = QPushButton("Обновить")
        refresh_button.clicked.connect(self.load_users)
        top_panel.addWidget(refresh_button)
        
        add_button = QPushButton("Добавить пользователя")
        add_button.clicked.connect(self.add_user)
        top_panel.addWidget(add_button)
        
        top_panel.addStretch()
        
        layout.addLayout(top_panel)
        
        # Таблица пользователей
        self.users_table = QTableWidget()
        self.users_table.setColumnCount(6)
        self.users_table.setHorizontalHeaderLabels([
            "Логин", "ФИО", "Роль", "Email", "Последний вход", "Действия"
        ])
        self.users_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        
        layout.addWidget(self.users_table)
        
        # Загрузка пользователей
        self.load_users()
    
    def load_users(self):
        """Загрузка списка пользователей"""
        users = db.get_all_users()
        
        self.users_table.setRowCount(len(users))
        
        # Маппинг ролей для отображения
        role_map = {
            'admin': 'Администратор',
            'doctor': 'Врач',
            'lab': 'Лаборант'
        }
        
        for row, user in enumerate(users):
            # Логин
            username_item = QTableWidgetItem(user.get('username', ''))
            username_item.setData(Qt.UserRole, user)  # Сохраняем все данные пользователя
            self.users_table.setItem(row, 0, username_item)
            
            # ФИО
            self.users_table.setItem(row, 1, QTableWidgetItem(user.get('full_name', '')))
            
            # Роль
            role_item = QTableWidgetItem(role_map.get(user.get('role', ''), user.get('role', '')))
            self.users_table.setItem(row, 2, role_item)
            
            # Email
            self.users_table.setItem(row, 3, QTableWidgetItem(user.get('email', '')))
            
            # Последний вход
            last_login = user.get('last_login', '')
            self.users_table.setItem(row, 4, QTableWidgetItem(last_login))
            
            # Ячейка с кнопками действий
            actions_widget = QWidget()
            actions_layout = QHBoxLayout(actions_widget)
            actions_layout.setContentsMargins(0, 0, 0, 0)
            
            edit_button = QPushButton("Редактировать")
            edit_button.clicked.connect(lambda checked, u=user: self.edit_user(u))
            actions_layout.addWidget(edit_button)
            
            delete_button = QPushButton("Удалить")
            delete_button.clicked.connect(lambda checked, u=user: self.delete_user(u))
            actions_layout.addWidget(delete_button)
            
            self.users_table.setCellWidget(row, 5, actions_widget)
    
    def add_user(self):
        """Добавление нового пользователя"""
        dialog = AddEditUserDialog(parent=self)
        if dialog.exec() == QDialog.Accepted:
            self.load_users()  # Обновляем список пользователей
    
    def edit_user(self, user):
        """Редактирование существующего пользователя"""
        dialog = AddEditUserDialog(user, parent=self)
        if dialog.exec() == QDialog.Accepted:
            self.load_users()  # Обновляем список пользователей
    
    def delete_user(self, user):
        """Удаление пользователя"""
        reply = QMessageBox.question(
            self,
            "Подтверждение удаления",
            f"Вы уверены, что хотите удалить пользователя {user.get('username')}?",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No
        )
        
        if reply == QMessageBox.Yes:
            try:
                # Удаление врача, если это врач
                if user.get('role') == 'doctor':
                    db.execute_query("DELETE FROM doctors WHERE user_id = ?", (user.get('id'),))
                
                # Удаление пользователя
                result = db.execute_query("DELETE FROM users WHERE id = ?", (user.get('id'),))
                
                if result:
                    QMessageBox.information(self, "Успех", "Пользователь успешно удален")
                    self.load_users()  # Обновляем список пользователей
                else:
                    QMessageBox.warning(self, "Ошибка", "Не удалось удалить пользователя")
            
            except Exception as e:
                QMessageBox.critical(self, "Ошибка", f"Произошла ошибка при удалении: {str(e)}")


class SystemStatisticsWidget(QWidget):
    """Виджет для отображения статистики системы"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setup_ui()
    
    def setup_ui(self):
        """Настройка интерфейса"""
        layout = QVBoxLayout(self)
        
        # Верхняя панель с фильтрами
        top_panel = QHBoxLayout()
        
        period_label = QLabel("Период:")
        self.start_date = QDateEdit()
        self.start_date.setDate(QDate.currentDate().addDays(-30))
        self.start_date.setCalendarPopup(True)
        
        to_label = QLabel("по")
        
        self.end_date = QDateEdit()
        self.end_date.setDate(QDate.currentDate())
        self.end_date.setCalendarPopup(True)
        
        apply_button = QPushButton("Применить")
        apply_button.clicked.connect(self.load_statistics)
        
        top_panel.addWidget(period_label)
        top_panel.addWidget(self.start_date)
        top_panel.addWidget(to_label)
        top_panel.addWidget(self.end_date)
        top_panel.addWidget(apply_button)
        top_panel.addStretch()
        
        # Кнопки для создания отчетов
        report_button = QPushButton("Создать отчет Excel")
        report_button.setStyleSheet("""
            QPushButton {
                background-color: #28a745;
                color: white;
                padding: 5px 10px;
                border-radius: 3px;
            }
            QPushButton:hover {
                background-color: #218838;
            }
        """)
        report_button.clicked.connect(self.generate_excel_report)
        top_panel.addWidget(report_button)
        
        csv_button = QPushButton("Создать отчет CSV")
        csv_button.setStyleSheet("""
            QPushButton {
                background-color: #17a2b8;
                color: white;
                padding: 5px 10px;
                border-radius: 3px;
            }
            QPushButton:hover {
                background-color: #138496;
            }
        """)
        csv_button.clicked.connect(self.generate_csv_report)
        top_panel.addWidget(csv_button)
        
        layout.addLayout(top_panel)
        
        # Панель с блоками статистики
        self.statistics_area = QScrollArea()
        self.statistics_area.setWidgetResizable(True)
        
        statistics_widget = QWidget()
        self.statistics_layout = QVBoxLayout(statistics_widget)
        
        self.statistics_area.setWidget(statistics_widget)
        layout.addWidget(self.statistics_area)
        
        # Загрузка статистики
        self.load_statistics()
    
    def load_statistics(self):
        """Загрузка статистики"""
        # Очистка предыдущих данных
        for i in reversed(range(self.statistics_layout.count())):
            widget = self.statistics_layout.itemAt(i).widget()
            if widget:
                widget.setParent(None)
        
        # Получение параметров фильтрации
        start_date = self.start_date.date().toString("yyyy-MM-dd")
        end_date = self.end_date.date().toString("yyyy-MM-dd")
        
        # Статистика пользователей
        self._add_user_statistics()
        
        # Статистика пациентов
        self._add_patient_statistics(start_date, end_date)
        
        # Статистика анализов
        self._add_analysis_statistics(start_date, end_date)
        
        # Статистика приемов
        self._add_appointment_statistics(start_date, end_date)
    
    def _add_user_statistics(self):
        """Добавление блока статистики пользователей"""
        group_box = QGroupBox("Статистика пользователей")
        layout = QVBoxLayout()
        
        # Общее количество пользователей
        users = db.fetch_all("SELECT role, COUNT(*) as count FROM users GROUP BY role")
        
        # Маппинг ролей для отображения
        role_map = {
            'admin': 'Администраторы',
            'doctor': 'Врачи',
            'lab': 'Лаборанты'
        }
        
        user_stats = QLabel("Пользователи в системе:")
        layout.addWidget(user_stats)
        
        for role_stat in users:
            role = role_stat.get('role', '')
            count = role_stat.get('count', 0)
            role_label = QLabel(f"• {role_map.get(role, role)}: {count}")
            layout.addWidget(role_label)
        
        group_box.setLayout(layout)
        self.statistics_layout.addWidget(group_box)
    
    def _add_patient_statistics(self, start_date, end_date):
        """Добавление блока статистики пациентов"""
        group_box = QGroupBox("Статистика пациентов")
        layout = QVBoxLayout()
        
        # Общее количество пациентов
        total_patients = db.fetch_one("SELECT COUNT(*) as count FROM patients")
        total_label = QLabel(f"Всего пациентов: {total_patients.get('count', 0)}")
        layout.addWidget(total_label)
        
        # Количество новых пациентов за период
        new_patients = db.fetch_one(
            "SELECT COUNT(*) as count FROM patients WHERE date(created_at) BETWEEN ? AND ?",
            (start_date, end_date)
        )
        new_patients_label = QLabel(f"Новых пациентов за период: {new_patients.get('count', 0)}")
        layout.addWidget(new_patients_label)
        
        group_box.setLayout(layout)
        self.statistics_layout.addWidget(group_box)
    
    def _add_analysis_statistics(self, start_date, end_date):
        """Добавление блока статистики анализов"""
        group_box = QGroupBox("Статистика анализов")
        layout = QVBoxLayout()
        
        # Общее количество анализов за период
        total_analyses = db.fetch_one(
            "SELECT COUNT(*) as count FROM analysis_results WHERE date(result_date) BETWEEN ? AND ?",
            (start_date, end_date)
        )
        total_label = QLabel(f"Всего анализов за период: {total_analyses.get('count', 0)}")
        layout.addWidget(total_label)
        
        # Количество анализов по типам
        analyses_by_type = db.fetch_all("""
            SELECT at.name, COUNT(ar.id) as count
            FROM analysis_results ar
            JOIN analysis_types at ON ar.analysis_type_id = at.id
            WHERE date(ar.result_date) BETWEEN ? AND ?
            GROUP BY at.name
            ORDER BY count DESC
        """, (start_date, end_date))
        
        if analyses_by_type:
            types_label = QLabel("Анализы по типам:")
            layout.addWidget(types_label)
            
            for type_stat in analyses_by_type:
                type_name = type_stat.get('name', '')
                count = type_stat.get('count', 0)
                type_label = QLabel(f"• {type_name}: {count}")
                layout.addWidget(type_label)
        else:
            no_data_label = QLabel("Нет данных о проведенных анализах за указанный период")
            layout.addWidget(no_data_label)
        
        group_box.setLayout(layout)
        self.statistics_layout.addWidget(group_box)
    
    def _add_appointment_statistics(self, start_date, end_date):
        """Добавление блока статистики приемов"""
        group_box = QGroupBox("Статистика приемов")
        layout = QVBoxLayout()
        
        # Общее количество приемов за период
        total_appointments = db.fetch_one(
            "SELECT COUNT(*) as count FROM appointments WHERE date(appointment_date) BETWEEN ? AND ?",
            (start_date, end_date)
        )
        total_label = QLabel(f"Всего приемов за период: {total_appointments.get('count', 0)}")
        layout.addWidget(total_label)
        
        # Количество приемов по статусам
        appointments_by_status = db.fetch_all("""
            SELECT status, COUNT(*) as count
            FROM appointments
            WHERE date(appointment_date) BETWEEN ? AND ?
            GROUP BY status
        """, (start_date, end_date))
        
        # Маппинг статусов для отображения
        status_map = {
            'scheduled': 'Запланировано',
            'completed': 'Завершено',
            'cancelled': 'Отменено'
        }
        
        if appointments_by_status:
            status_label = QLabel("Приемы по статусам:")
            layout.addWidget(status_label)
            
            for status_stat in appointments_by_status:
                status = status_stat.get('status', '')
                count = status_stat.get('count', 0)
                status_label = QLabel(f"• {status_map.get(status, status)}: {count}")
                layout.addWidget(status_label)
        else:
            no_data_label = QLabel("Нет данных о приемах за указанный период")
            layout.addWidget(no_data_label)
        
        group_box.setLayout(layout)
        self.statistics_layout.addWidget(group_box)
    
    def generate_excel_report(self):
        """Создание отчета в формате Excel"""
        try:
            import xlwt
            from datetime import datetime
            
            # Создаем книгу Excel
            wb = xlwt.Workbook()
            
            # Создаем лист для общей статистики
            ws_general = wb.add_sheet('Общая статистика')
            
            # Заголовки
            ws_general.write(0, 0, "Медицинский центр - Статистика")
            ws_general.write(1, 0, f"Период: {self.start_date.date().toString('dd.MM.yyyy')} - {self.end_date.date().toString('dd.MM.yyyy')}")
            ws_general.write(2, 0, f"Дата создания отчета: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
            
            # Получаем данные
            start_date = self.start_date.date().toString("yyyy-MM-dd")
            end_date = self.end_date.date().toString("yyyy-MM-dd")
            
            # Статистика пользователей
            row = 4
            ws_general.write(row, 0, "Статистика пользователей")
            row += 1
            
            users = db.fetch_all("SELECT role, COUNT(*) as count FROM users GROUP BY role")
            role_map = {
                'admin': 'Администраторы',
                'doctor': 'Врачи',
                'lab': 'Лаборанты'
            }
            
            for role_stat in users:
                role = role_stat.get('role', '')
                count = role_stat.get('count', 0)
                ws_general.write(row, 0, f"{role_map.get(role, role)}")
                ws_general.write(row, 1, count)
                row += 1
            
            # Статистика пациентов
            row += 1
            ws_general.write(row, 0, "Статистика пациентов")
            row += 1
            
            total_patients = db.fetch_one("SELECT COUNT(*) as count FROM patients")
            ws_general.write(row, 0, "Всего пациентов")
            ws_general.write(row, 1, total_patients.get('count', 0))
            row += 1
            
            new_patients = db.fetch_one(
                "SELECT COUNT(*) as count FROM patients WHERE date(created_at) BETWEEN ? AND ?",
                (start_date, end_date)
            )
            ws_general.write(row, 0, "Новых пациентов за период")
            ws_general.write(row, 1, new_patients.get('count', 0))
            row += 1
            
            # Статистика анализов
            row += 1
            ws_general.write(row, 0, "Статистика анализов")
            row += 1
            
            analyses_by_type = db.fetch_all("""
                SELECT at.name, COUNT(ar.id) as count
                FROM analysis_results ar
                JOIN analysis_types at ON ar.analysis_type_id = at.id
                WHERE date(ar.result_date) BETWEEN ? AND ?
                GROUP BY at.name
                ORDER BY count DESC
            """, (start_date, end_date))
            
            for type_stat in analyses_by_type:
                type_name = type_stat.get('name', '')
                count = type_stat.get('count', 0)
                ws_general.write(row, 0, type_name)
                ws_general.write(row, 1, count)
                row += 1
            
            # Создаем лист со списком пациентов
            ws_patients = wb.add_sheet('Список пациентов')
            
            # Заголовки столбцов
            headers = ["ID", "ФИО", "Дата рождения", "Телефон", "Email", "Адрес"]
            for col, header in enumerate(headers):
                ws_patients.write(0, col, header)
            
            # Данные пациентов
            patients = db.get_all_patients()
            for row, patient in enumerate(patients, 1):
                ws_patients.write(row, 0, patient.get('id', ''))
                ws_patients.write(row, 1, patient.get('full_name', ''))
                ws_patients.write(row, 2, patient.get('birth_date', ''))
                ws_patients.write(row, 3, patient.get('phone', ''))
                ws_patients.write(row, 4, patient.get('email', ''))
                ws_patients.write(row, 5, patient.get('address', ''))
            
            # Сохраняем отчет
            report_name = f"medical_report_Общий_список_пациентов_{datetime.now().strftime('%Y-%m-%d')}.xls"
            wb.save(report_name)
            
            QMessageBox.information(self, "Отчет создан", f"Отчет успешно сохранен: {report_name}")
        
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка при создании отчета: {str(e)}")
    
    def generate_csv_report(self):
        """Создание отчета в формате CSV"""
        try:
            import csv
            from datetime import datetime
            
            # Получаем данные
            patients = db.get_all_patients()
            
            # Создаем CSV-файл
            report_name = f"medical_report_Общий_список_пациентов_{datetime.now().strftime('%Y-%m-%d')}.csv"
            
            with open(report_name, 'w', newline='', encoding='utf-8') as csvfile:
                writer = csv.writer(csvfile)
                
                # Заголовок отчета
                writer.writerow(["Медицинский центр - Список пациентов"])
                writer.writerow([f"Дата создания: {datetime.now().strftime('%d.%m.%Y %H:%M')}"])
                writer.writerow([])  # Пустая строка
                
                # Заголовки столбцов
                writer.writerow(["ID", "ФИО", "Дата рождения", "Телефон", "Email", "Адрес"])
                
                # Данные пациентов
                for patient in patients:
                    writer.writerow([
                        patient.get('id', ''),
                        patient.get('full_name', ''),
                        patient.get('birth_date', ''),
                        patient.get('phone', ''),
                        patient.get('email', ''),
                        patient.get('address', '')
                    ])
            
            QMessageBox.information(self, "Отчет создан", f"Отчет CSV успешно сохранен: {report_name}")
        
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка при создании отчета CSV: {str(e)}")


class PatientListWidget(QWidget):
    """Виджет для отображения списка пациентов"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setup_ui()
    
    def setup_ui(self):
        """Настройка интерфейса"""
        layout = QVBoxLayout(self)
        
        # Верхняя панель с кнопками и поиском
        top_panel = QHBoxLayout()
        
        # Кнопка добавления нового пациента
        add_button = QPushButton("Добавить пациента")
        add_button.setStyleSheet("""
            QPushButton {
                background-color: #28a745;
                color: white;
                padding: 5px 10px;
                border-radius: 3px;
            }
            QPushButton:hover {
                background-color: #218838;
            }
        """)
        add_button.clicked.connect(self.add_patient)
        top_panel.addWidget(add_button)
        
        # Поле поиска
        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText("Поиск пациента")
        self.search_input.textChanged.connect(self.filter_patients)
        top_panel.addWidget(self.search_input)
        
        # Кнопка обновления списка
        refresh_button = QPushButton("Обновить")
        refresh_button.clicked.connect(self.load_patients)
        top_panel.addWidget(refresh_button)
        
        top_panel.addStretch()
        
        layout.addLayout(top_panel)
        
        # Таблица пациентов
        self.patients_table = QTableWidget()
        self.patients_table.setColumnCount(7)
        self.patients_table.setHorizontalHeaderLabels([
            "ID", "ФИО", "Дата рождения", "Пол", "Телефон", "Email", "Действия"
        ])
        self.patients_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.patients_table.setSortingEnabled(True)
        
        layout.addWidget(self.patients_table)
        
        # Загрузка пациентов
        self.load_patients()
    
    def load_patients(self):
        """Загрузка списка пациентов"""
        patients = db.get_all_patients()
        self.all_patients = patients  # Сохраняем для фильтрации
        
        self.update_table(patients)
    
    def update_table(self, patients):
        """Обновление таблицы пациентов"""
        self.patients_table.setRowCount(len(patients))
        
        for row, patient in enumerate(patients):
            # ID
            self.patients_table.setItem(row, 0, QTableWidgetItem(str(patient.get('id', ''))))
            
            # ФИО
            self.patients_table.setItem(row, 1, QTableWidgetItem(patient.get('full_name', '')))
            
            # Дата рождения
            self.patients_table.setItem(row, 2, QTableWidgetItem(patient.get('birth_date', '')))
            
            # Пол (может отсутствовать в базе)
            gender = patient.get('gender', 'Не указан')
            self.patients_table.setItem(row, 3, QTableWidgetItem(gender))
            
            # Телефон
            self.patients_table.setItem(row, 4, QTableWidgetItem(patient.get('phone', '')))
            
            # Email
            self.patients_table.setItem(row, 5, QTableWidgetItem(patient.get('email', '')))
            
            # Кнопки действий
            actions_widget = QWidget()
            actions_layout = QHBoxLayout(actions_widget)
            actions_layout.setContentsMargins(0, 0, 0, 0)
            
            edit_button = QPushButton("Редактировать")
            edit_button.setStyleSheet("min-width: 80px;")
            edit_button.clicked.connect(lambda checked, p=patient: self.edit_patient(p))
            actions_layout.addWidget(edit_button)
            
            add_appointment_button = QPushButton("Запись")
            add_appointment_button.setStyleSheet("min-width: 80px;")
            add_appointment_button.clicked.connect(lambda checked, p=patient: self.add_appointment(p))
            actions_layout.addWidget(add_appointment_button)
            
            delete_button = QPushButton("Удалить")
            delete_button.setStyleSheet("min-width: 80px;")
            delete_button.clicked.connect(lambda checked, p=patient: self.delete_patient(p))
            actions_layout.addWidget(delete_button)
            
            self.patients_table.setCellWidget(row, 6, actions_widget)
    
    def filter_patients(self):
        """Фильтрация пациентов по поисковому запросу"""
        search_text = self.search_input.text().lower()
        
        if not search_text:
            # Если строка поиска пуста, показываем всех пациентов
            self.update_table(self.all_patients)
            return
        
        # Фильтрация пациентов
        filtered_patients = []
        for patient in self.all_patients:
            full_name = patient.get('full_name', '').lower()
            phone = patient.get('phone', '').lower()
            email = patient.get('email', '').lower()
            
            if (search_text in full_name or 
                search_text in phone or 
                search_text in email):
                filtered_patients.append(patient)
        
        self.update_table(filtered_patients)
    
    def add_patient(self):
        """Добавление нового пациента"""
        dialog = AddPatientDialog(parent=self)
        if dialog.exec() == QDialog.Accepted:
            self.load_patients()  # Обновляем список пациентов
    
    def edit_patient(self, patient):
        """Редактирование пациента"""
        dialog = AddPatientDialog(patient=patient, parent=self)
        if dialog.exec() == QDialog.Accepted:
            self.load_patients()  # Обновляем список пациентов
    
    def delete_patient(self, patient):
        """Удаление пациента"""
        reply = QMessageBox.question(
            self,
            "Подтверждение удаления",
            f"Вы уверены, что хотите удалить пациента {patient.get('full_name')}?",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No
        )
        
        if reply == QMessageBox.Yes:
            try:
                result = db.execute_query("DELETE FROM patients WHERE id = ?", (patient.get('id'),))
                
                if result:
                    QMessageBox.information(self, "Успех", "Пациент успешно удален")
                    self.load_patients()  # Обновляем список пациентов
                else:
                    QMessageBox.warning(self, "Ошибка", "Не удалось удалить пациента")
            
            except Exception as e:
                QMessageBox.critical(self, "Ошибка", f"Произошла ошибка при удалении: {str(e)}")
    
    def add_appointment(self, patient):
        """Запись пациента на прием"""
        try:
            # Получаем доступ к главному окну
            main_window = None
            
            # Проверяем, является ли непосредственный родитель главным окном
            if hasattr(self.parent(), 'add_appointment_dialog'):
                main_window = self.parent()
            else:
                # Если нет, ищем главное окно в иерархии
                parent = self.parent()
                while parent:
                    if hasattr(parent, 'add_appointment_dialog'):
                        main_window = parent
                        break
                    parent = parent.parent()
                
                # Если и это не помогло, пробуем получить главное окно через метод window()
                if not main_window:
                    main_window = self.window()
            
            # Если нашли главное окно и у него есть нужный метод
            if main_window and hasattr(main_window, 'add_appointment_dialog'):
                print(f"Вызываем метод add_appointment_dialog для пациента: {patient.get('full_name')}")
                main_window.add_appointment_dialog(patient)
            else:
                # Если не нашли, выводим сообщение об ошибке
                print("Не удалось найти метод add_appointment_dialog в родительском окне")
                QMessageBox.warning(
                    self, 
                    "Ошибка", 
                    "Не удалось открыть форму записи на прием. Пожалуйста, используйте вкладку 'Записи на прием'."
                )
        
        except Exception as e:
            print(f"Ошибка при создании записи на прием: {str(e)}")
            QMessageBox.warning(
                self, 
                "Ошибка", 
                f"Произошла ошибка при создании записи: {str(e)}"
            )


class AddPatientDialog(QDialog):
    """Диалоговое окно для добавления/редактирования пациента"""
    
    def __init__(self, patient=None, parent=None):
        super().__init__(parent)
        self.patient = patient
        
        title = "Редактирование пациента" if patient else "Добавление пациента"
        self.setWindowTitle(title)
        self.setMinimumWidth(450)
        self.setup_ui()
    
    def setup_ui(self):
        """Настройка пользовательского интерфейса"""
        layout = QVBoxLayout()
        
        # Форма с полями ввода
        form_layout = QFormLayout()
        
        # ФИО
        self.full_name_input = QLineEdit()
        if self.patient:
            self.full_name_input.setText(self.patient.get('full_name', ''))
        form_layout.addRow("ФИО:", self.full_name_input)
        
        # Дата рождения
        self.birth_date_input = QDateEdit()
        self.birth_date_input.setCalendarPopup(True)
        if self.patient and self.patient.get('birth_date'):
            try:
                date = QDate.fromString(self.patient.get('birth_date'), "yyyy-MM-dd")
                self.birth_date_input.setDate(date)
            except:
                self.birth_date_input.setDate(QDate.currentDate())
        else:
            self.birth_date_input.setDate(QDate.currentDate())
        form_layout.addRow("Дата рождения:", self.birth_date_input)
        
        # Пол
        gender_layout = QHBoxLayout()
        self.gender_group = QButtonGroup(self)
        
        self.male_radio = QRadioButton("Мужской")
        self.female_radio = QRadioButton("Женский")
        
        if self.patient:
            if self.patient.get('gender') == "Женский":
                self.female_radio.setChecked(True)
            else:
                self.male_radio.setChecked(True)
        else:
            self.male_radio.setChecked(True)
        
        self.gender_group.addButton(self.male_radio)
        self.gender_group.addButton(self.female_radio)
        
        gender_layout.addWidget(self.male_radio)
        gender_layout.addWidget(self.female_radio)
        gender_layout.addStretch()
        
        form_layout.addRow("Пол:", gender_layout)
        
        # Телефон
        self.phone_input = QLineEdit()
        if self.patient:
            self.phone_input.setText(self.patient.get('phone', ''))
        form_layout.addRow("Телефон:", self.phone_input)
        
        # Email
        self.email_input = QLineEdit()
        if self.patient:
            self.email_input.setText(self.patient.get('email', ''))
        form_layout.addRow("Email:", self.email_input)
        
        # Адрес
        self.address_input = QLineEdit()
        if self.patient:
            self.address_input.setText(self.patient.get('address', ''))
        form_layout.addRow("Адрес:", self.address_input)
        
        layout.addLayout(form_layout)
        
        # Кнопки
        buttons_layout = QHBoxLayout()
        
        save_button = QPushButton("Сохранить")
        save_button.clicked.connect(self.save_patient)
        buttons_layout.addWidget(save_button)
        
        cancel_button = QPushButton("Отмена")
        cancel_button.clicked.connect(self.reject)
        buttons_layout.addWidget(cancel_button)
        
        layout.addLayout(buttons_layout)
        
        self.setLayout(layout)
    
    def save_patient(self):
        """Сохранение данных пациента"""
        # Получение данных из формы
        full_name = self.full_name_input.text().strip()
        birth_date = self.birth_date_input.date().toString("yyyy-MM-dd")
        gender = "Женский" if self.female_radio.isChecked() else "Мужской"
        phone = self.phone_input.text().strip()
        email = self.email_input.text().strip()
        address = self.address_input.text().strip()
        
        # Валидация данных
        if not full_name:
            QMessageBox.warning(self, "Ошибка", "ФИО пациента не может быть пустым")
            return
        
        try:
            if self.patient:  # Редактирование
                # Проверяем наличие столбца gender
                try:
                    columns_query = "PRAGMA table_info(patients)"
                    columns = db.fetch_all(columns_query)
                    column_names = [col.get('name') for col in columns]
                    
                    if 'gender' in column_names:
                        # Если столбец gender существует
                        success = db.execute_query(
                            """UPDATE patients 
                               SET full_name = ?, birth_date = ?, gender = ?, phone = ?, email = ?, address = ? 
                               WHERE id = ?""",
                            (full_name, birth_date, gender, phone, email, address, self.patient.get('id'))
                        )
                    else:
                        # Если столбца gender нет
                        success = db.execute_query(
                            """UPDATE patients 
                               SET full_name = ?, birth_date = ?, phone = ?, email = ?, address = ? 
                               WHERE id = ?""",
                            (full_name, birth_date, phone, email, address, self.patient.get('id'))
                        )
                    
                    if success:
                        QMessageBox.information(self, "Успех", "Пациент успешно обновлен")
                        self.accept()
                    else:
                        print("Не удалось обновить данные пациента, но ошибка не возникла")
                        QMessageBox.warning(self, "Примечание", "Возникла проблема при обновлении данных, но пациент мог быть обновлен")
                        self.accept()  # Все равно закрываем диалог
                
                except Exception as db_error:
                    print(f"Ошибка при работе с базой данных: {str(db_error)}")
                    
                    # Пробуем упрощенный запрос без gender
                    try:
                        simple_success = db.execute_query(
                            """UPDATE patients 
                               SET full_name = ?, birth_date = ?, phone = ?, email = ?, address = ? 
                               WHERE id = ?""",
                            (full_name, birth_date, phone, email, address, self.patient.get('id'))
                        )
                        
                        if simple_success:
                            QMessageBox.information(self, "Успех", "Пациент успешно обновлен (упрощенный режим)")
                            self.accept()
                        else:
                            QMessageBox.warning(self, "Ошибка", "Не удалось обновить данные пациента")
                    
                    except Exception as e:
                        print(f"Ошибка при упрощенном обновлении: {str(e)}")
                        QMessageBox.critical(self, "Ошибка", f"Не удалось обновить данные: {str(e)}")
            
            else:  # Добавление
                try:
                    # Проверяем наличие столбца gender
                    columns_query = "PRAGMA table_info(patients)"
                    columns = db.fetch_all(columns_query)
                    column_names = [col.get('name') for col in columns]
                    
                    if 'gender' in column_names:
                        # Если столбец gender существует
                        success = db.execute_query(
                            """INSERT INTO patients (full_name, birth_date, gender, phone, email, address) 
                               VALUES (?, ?, ?, ?, ?, ?)""",
                            (full_name, birth_date, gender, phone, email, address)
                        )
                    else:
                        # Если столбца gender нет
                        success = db.execute_query(
                            """INSERT INTO patients (full_name, birth_date, phone, email, address) 
                               VALUES (?, ?, ?, ?, ?)""",
                            (full_name, birth_date, phone, email, address)
                        )
                    
                    if success:
                        QMessageBox.information(self, "Успех", "Пациент успешно добавлен")
                        self.accept()
                    else:
                        print("Не удалось добавить пациента, но ошибка не возникла")
                        QMessageBox.warning(self, "Примечание", "Возникла проблема при добавлении, но пациент мог быть добавлен")
                        self.accept()  # Все равно закрываем диалог
                
                except Exception as db_error:
                    print(f"Ошибка при работе с базой данных: {str(db_error)}")
                    
                    # Пробуем упрощенный запрос без gender
                    try:
                        simple_success = db.execute_query(
                            """INSERT INTO patients (full_name, birth_date, phone, email, address) 
                               VALUES (?, ?, ?, ?, ?)""",
                            (full_name, birth_date, phone, email, address)
                        )
                        
                        if simple_success:
                            QMessageBox.information(self, "Успех", "Пациент успешно добавлен (упрощенный режим)")
                            self.accept()
                        else:
                            QMessageBox.warning(self, "Ошибка", "Не удалось добавить пациента")
                    
                    except Exception as e:
                        print(f"Ошибка при упрощенном добавлении: {str(e)}")
                        QMessageBox.critical(self, "Ошибка", f"Не удалось добавить пациента: {str(e)}")
        
        except Exception as e:
            print(f"Общая ошибка при сохранении пациента: {str(e)}")
            QMessageBox.critical(self, "Ошибка", f"Произошла ошибка: {str(e)}")


class AnalysisResultsWidget(QWidget):
    """Виджет для работы с результатами анализов"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setup_ui()
        self.document_generator = DocumentGenerator()
        self.email_sender = EmailSender(test_mode=False)
    
    def setup_ui(self):
        """Настройка интерфейса"""
        layout = QVBoxLayout(self)
        
        # Фильтры для анализов
        filters_group = QGroupBox("Фильтры")
        filters_layout = QHBoxLayout()
        
        # Выбор пациента
        patient_label = QLabel("Пациент:")
        self.patient_combo = QComboBox()
        self.patient_combo.setMinimumWidth(200)
        
        # Добавляем опцию "Все пациенты"
        self.patient_combo.addItem("Все пациенты", None)
        
        # Загружаем список пациентов
        patients = db.get_all_patients()
        for patient in patients:
            self.patient_combo.addItem(patient.get('full_name', ''), patient.get('id'))
        
        # Выбор типа анализа
        analysis_type_label = QLabel("Тип анализа:")
        self.analysis_type_combo = QComboBox()
        self.analysis_type_combo.setMinimumWidth(150)
        
        # Добавляем опцию "Все типы"
        self.analysis_type_combo.addItem("Все типы", None)
        
        # Загружаем типы анализов
        analysis_types = db.fetch_all("SELECT * FROM analysis_types")
        for analysis_type in analysis_types:
            self.analysis_type_combo.addItem(analysis_type.get('name', ''), analysis_type.get('id'))
        
        # Выбор периода дат
        date_from_label = QLabel("С:")
        self.date_from = QDateEdit()
        self.date_from.setDate(QDate.currentDate().addDays(-30))
        self.date_from.setCalendarPopup(True)
        
        date_to_label = QLabel("По:")
        self.date_to = QDateEdit()
        self.date_to.setDate(QDate.currentDate())
        self.date_to.setCalendarPopup(True)
        
        # Кнопки управления фильтрами
        apply_filters_button = QPushButton("Применить")
        apply_filters_button.clicked.connect(self.refresh_analysis_results)
        
        clear_filters_button = QPushButton("Сбросить")
        clear_filters_button.clicked.connect(self.clear_filters)
        
        # Размещение элементов фильтра
        filters_layout.addWidget(patient_label)
        filters_layout.addWidget(self.patient_combo)
        filters_layout.addWidget(analysis_type_label)
        filters_layout.addWidget(self.analysis_type_combo)
        filters_layout.addWidget(date_from_label)
        filters_layout.addWidget(self.date_from)
        filters_layout.addWidget(date_to_label)
        filters_layout.addWidget(self.date_to)
        filters_layout.addWidget(apply_filters_button)
        filters_layout.addWidget(clear_filters_button)
        
        filters_group.setLayout(filters_layout)
        layout.addWidget(filters_group)
        
        # Таблица результатов анализов
        self.results_table = QTableWidget()
        self.results_table.setColumnCount(7)
        self.results_table.setHorizontalHeaderLabels([
            "Дата", "Пациент", "Тип анализа", "Статус", "Лаборант", "Документы", "Email"
        ])
        self.results_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.results_table.setSortingEnabled(True)
        
        layout.addWidget(self.results_table)
        
        # Кнопки действий
        actions_layout = QHBoxLayout()
        
        # Кнопка экспорта в Excel
        export_excel_button = QPushButton("Экспорт в Excel")
        export_excel_button.clicked.connect(self.export_to_excel)
        actions_layout.addWidget(export_excel_button)
        
        # Кнопка экспорта всех результатов в Word
        export_all_word_button = QPushButton("Экспорт всех в Word")
        export_all_word_button.clicked.connect(self.export_all_to_word)
        actions_layout.addWidget(export_all_word_button)
        
        # Кнопка отправки отчета по email
        send_report_button = QPushButton("Отправить отчет")
        send_report_button.clicked.connect(self.send_report_by_email)
        actions_layout.addWidget(send_report_button)
        
        layout.addLayout(actions_layout)
        
        # Загрузка результатов анализов
        self.refresh_analysis_results()
    
    def refresh_analysis_results(self):
        """Обновление списка результатов анализов с учетом фильтров"""
        # Получение параметров фильтрации
        patient_id = self.patient_combo.currentData()
        analysis_type_id = self.analysis_type_combo.currentData()
        from_date = self.date_from.date().toString("yyyy-MM-dd")
        to_date = self.date_to.date().toString("yyyy-MM-dd")
        
        # Запрос к базе данных
        # Составление условий запроса
        conditions = ["1=1"]  # Всегда истинное условие для удобства добавления других через AND
        params = []
        
        if patient_id:
            conditions.append("ar.patient_id = ?")
            params.append(patient_id)
        
        if analysis_type_id:
            conditions.append("ar.analysis_type_id = ?")
            params.append(analysis_type_id)
        
        conditions.append("date(ar.result_date) BETWEEN ? AND ?")
        params.extend([from_date, to_date])
        
        # Формирование и выполнение запроса
        query = f"""
            SELECT ar.id as id, p.full_name as patient_name, p.birth_date, at.name as analysis_type, 
                   ar.result_date, u.full_name as lab_technician, ar.status
            FROM analysis_results ar
            JOIN patients p ON ar.patient_id = p.id
            JOIN analysis_types at ON ar.analysis_type_id = at.id
            JOIN users u ON ar.lab_user_id = u.id
            WHERE {' AND '.join(conditions)}
            ORDER BY ar.result_date DESC
        """
        results = db.fetch_all(query, params)
        
        # Отладочный вывод результатов запроса
        print(f"Получено {len(results)} результатов анализов")
        
        # Обновление таблицы
        self.results_table.setRowCount(len(results))
        
        for row, result in enumerate(results):
            # Проверяем, содержит ли результат ID
            if 'id' not in result:
                print(f"ОШИБКА: Запись не содержит ID для строки {row}")
                continue
            
            # Распаковка данных
            result_id = result.get('id')
            
            # Проверка, что ID существует и является числом
            if result_id is None:
                print(f"ID результата отсутствует в данных")
                continue
            
            # Гарантируем, что ID - целое число
            try:
                result_id = int(result_id)
            except (ValueError, TypeError):
                print(f"Ошибка преобразования ID {result_id} в число")
                continue
            
            patient_name = result.get('patient_name', '')
            birth_date = result.get('birth_date', '')
            analysis_type = result.get('analysis_type', '')
            result_date = result.get('result_date', '')
            lab_technician = result.get('lab_technician', '')
            status = result.get('status', '')
            
            # Установка данных в ячейки таблицы
            self.results_table.setItem(row, 0, QTableWidgetItem(result_date))
            self.results_table.setItem(row, 1, QTableWidgetItem(f"{patient_name} ({birth_date})"))
            self.results_table.setItem(row, 2, QTableWidgetItem(analysis_type))
            self.results_table.setItem(row, 3, QTableWidgetItem(self.translate_status(status)))
            self.results_table.setItem(row, 4, QTableWidgetItem(lab_technician))
            
            # Кнопки для работы с документами
            doc_widget = QWidget()
            doc_layout = QHBoxLayout(doc_widget)
            doc_layout.setContentsMargins(0, 0, 0, 0)
            
            view_button = QPushButton("Просмотр")
            view_button.setProperty("result_id", result_id)
            view_button.clicked.connect(self.view_analysis_result)
            
            doc_layout.addWidget(view_button)
            
            self.results_table.setCellWidget(row, 5, doc_widget)
            
            # Кнопка отправки по email
            email_button = QPushButton("Отправить")
            email_button.setProperty("result_id", result_id)
            email_button.clicked.connect(self.send_by_email)
            
            self.results_table.setCellWidget(row, 6, email_button)
    
    def clear_filters(self):
        """Сброс фильтров"""
        self.patient_combo.setCurrentIndex(0)  # "Все пациенты"
        self.analysis_type_combo.setCurrentIndex(0)  # "Все типы"
        self.date_from.setDate(QDate.currentDate().addDays(-30))
        self.date_to.setDate(QDate.currentDate())
        
        # Обновление списка результатов
        self.refresh_analysis_results()
    
    def translate_status(self, status):
        """Перевод статуса на русский язык"""
        status_map = {
            'pending': 'В обработке',
            'completed': 'Выполнен',
            'cancelled': 'Отменен',
            'sent': 'Отправлен'
        }
        return status_map.get(status, status)
    
    def view_analysis_result(self):
        """Просмотр результата анализа"""
        # Получение ID результата
        sender = self.sender()
        result_id = sender.property("result_id")
        
        # Проверяем, что ID является корректным
        try:
            result_id = int(result_id) if result_id is not None else 0
            if result_id <= 0:
                error_dialog = ErrorDialog(
                    self,
                    "Не удалось загрузить результат анализа.",
                    "Ошибка"
                )
                error_dialog.exec()
                return
        except (ValueError, TypeError):
            error_dialog = ErrorDialog(
                self,
                "Ошибка при обработке данных анализа.",
                "Ошибка"
            )
            error_dialog.exec()
            return
        
        # Загрузка деталей результата
        result_details = db.get_analysis_result_details(result_id)
        
        if result_details:
            # Создание диалога с результатами
            dialog = QDialog(self)
            dialog.setWindowTitle(f"Результаты анализа: {result_details['analysis_type']['name']}")
            dialog.setMinimumWidth(500)
            
            # Макет диалога
            layout = QVBoxLayout(dialog)
            
            # Информация о пациенте и анализе
            info_layout = QFormLayout()
            
            patient_label = QLabel(result_details['patient']['full_name'])
            patient_label.setFont(QFont("Arial", 10, QFont.Bold))
            info_layout.addRow("Пациент:", patient_label)
            
            info_layout.addRow("Дата рождения:", QLabel(result_details['patient']['birth_date']))
            info_layout.addRow("Тип анализа:", QLabel(result_details['analysis_type']['name']))
            info_layout.addRow("Дата взятия:", QLabel(result_details['date_taken']))
            info_layout.addRow("Лаборант:", QLabel(result_details['lab_technician']))
            
            layout.addLayout(info_layout)
            
            # Таблица с результатами
            results_table = QTableWidget()
            results_table.setColumnCount(4)
            results_table.setHorizontalHeaderLabels(["Параметр", "Значение", "Норма", "Оценка"])
            results_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
            
            layout.addWidget(QLabel("Результаты:"))
            layout.addWidget(results_table)
            
            # Заполнение таблицы результатов
            parameters = result_details.get('parameters', [])
            results_table.setRowCount(len(parameters))
            
            for row, param in enumerate(parameters):
                # Название параметра
                results_table.setItem(row, 0, QTableWidgetItem(param.get('name', '')))
                
                # Значение с единицами измерения
                value_text = f"{param.get('value', '')} {param.get('unit', '')}"
                results_table.setItem(row, 1, QTableWidgetItem(value_text))
                
                # Нормальный диапазон
                if param.get('normal_min') is not None and param.get('normal_max') is not None:
                    normal_range = f"{param.get('normal_min')} - {param.get('normal_max')} {param.get('unit', '')}"
                else:
                    normal_range = "Не определена"
                results_table.setItem(row, 2, QTableWidgetItem(normal_range))
                
                # Оценка (в норме или нет)
                is_normal = param.get('is_normal')
                if is_normal is True:
                    status_item = QTableWidgetItem("В норме")
                    status_item.setForeground(QColor(0, 128, 0))  # Зеленый цвет
                elif is_normal is False:
                    status_item = QTableWidgetItem("Отклонение")
                    status_item.setForeground(QColor(255, 0, 0))  # Красный цвет
                else:
                    status_item = QTableWidgetItem("Не определено")
                
                results_table.setItem(row, 3, status_item)
            
            # Кнопки действий
            buttons_layout = QHBoxLayout()
            
            # Кнопка экспорта в Word
            word_button = QPushButton("Экспорт в Word")
            word_button.clicked.connect(lambda: report_generator.export_analysis_to_word(result_id, dialog))
            buttons_layout.addWidget(word_button)
            
            # Кнопка отправки по email
            email_button = QPushButton("Отправить по Email")
            email_button.clicked.connect(lambda: self.send_by_email(result_id=result_id, dialog=dialog))
            buttons_layout.addWidget(email_button)
            
            # Кнопка закрытия
            close_button = QPushButton("Закрыть")
            close_button.clicked.connect(dialog.accept)
            buttons_layout.addWidget(close_button)
            
            layout.addLayout(buttons_layout)
            
            # Показываем диалог
            dialog.exec()
        else:
            error_dialog = ErrorDialog(
                self,
                "Не удалось загрузить результаты анализа.",
                "Ошибка"
            )
            error_dialog.exec()
    
    def export_to_word(self, result_id=None, dialog=None):
        """Экспорт результата анализа в Word"""
        # Получение ID результата, если он не был передан
        if result_id is None:
            sender = self.sender()
            result_id = sender.property("result_id")
            print(f"Получен ID из кнопки: {result_id}, тип: {type(result_id)}")
        
        # Проверка, что ID является числом и больше 0
        try:
            result_id = int(result_id) if result_id is not None else 0
            print(f"Преобразован ID в число: {result_id}")
            if result_id <= 0:
                print(f"Некорректный ID: {result_id} (должен быть > 0)")
                # Используем новый класс диалога вместо QMessageBox
                error_dialog = ErrorDialog(
                    dialog or self,
                    "Ошибка при загрузке результата анализа.",
                    "Ошибка"
                )
                error_dialog.exec()
                return
                
            # Дополнительная проверка существования записи
            check_query = "SELECT id FROM analysis_results WHERE id = ?"
            print(f"Проверка существования записи: {check_query} с ID={result_id}")
            result_exists = db.fetch_one(check_query, (result_id,))
            
            if not result_exists:
                print(f"Результат с ID {result_id} не найден в БД")
                error_dialog = ErrorDialog(
                    dialog or self,
                    f"Результат анализа не найден в базе данных",
                    "Ошибка"
                )
                error_dialog.exec()
                return
                
            print(f"Результат существует: {result_exists}")
            
        except (ValueError, TypeError) as e:
            print(f"Ошибка при преобразовании ID: {e}")
            error_dialog = ErrorDialog(
                dialog or self,
                "Ошибка при обработке данных анализа",
                "Ошибка"
            )
            error_dialog.exec()
            return
            
        # Используем функцию из модуля report_generator
        report_generator.export_analysis_to_word(result_id, dialog or self)
    
    def export_all_to_word(self):
        """Экспорт всех результатов анализов в Word"""
        # Получаем текущие фильтры из интерфейса
        filters = {}
        
        # Получаем ID пациента, если выбран
        if self.patient_combo.currentData():
            filters['patient_id'] = self.patient_combo.currentData()
        
        # Получаем ID типа анализа, если выбран
        if self.analysis_type_combo.currentData():
            filters['analysis_type_id'] = self.analysis_type_combo.currentData()
        
        # Получаем даты из полей выбора даты
        filters['from_date'] = self.date_from.date().toString("yyyy-MM-dd")
        filters['to_date'] = self.date_to.date().toString("yyyy-MM-dd")
        
        # Используем функцию из модуля report_generator
        report_generator.export_all_analyses_to_word(self, filters)
    
    def send_by_email(self, result_id=None, dialog=None):
        """Отправка результата анализа по email"""
        # Получение ID результата, если он не был передан
        if result_id is None:
            sender = self.sender()
            result_id = sender.property("result_id")
        
        # Проверяем, что ID является корректным
        try:
            result_id = int(result_id) if result_id is not None else 0
            if result_id <= 0:
                error_dialog = ErrorDialog(
                    dialog or self,
                    "Не удалось определить результат анализа для отправки.",
                    "Ошибка"
                )
                error_dialog.exec()
                return
        except (ValueError, TypeError):
            error_dialog = ErrorDialog(
                dialog or self,
                "Ошибка при обработке данных анализа.",
                "Ошибка"
            )
            error_dialog.exec()
            return
            
        # Загрузка деталей результата
        result_details = db.get_analysis_result_details(result_id)
        
        if not result_details:
            error_dialog = ErrorDialog(
                dialog or self,
                "Не удалось загрузить результаты анализа.",
                "Ошибка"
            )
            error_dialog.exec()
            return
        
        # Получение email пациента
        patient_email = result_details['patient'].get('email')
        
        if not patient_email:
            # Спрашиваем email, если он не указан
            email_dialog = QDialog(dialog or self)
            email_dialog.setWindowTitle("Укажите email для отправки")
            email_dialog.setMinimumWidth(300)
            
            layout = QVBoxLayout(email_dialog)
            
            email_label = QLabel(f"Email для отправки результатов анализа {result_details['analysis_type']['name']} пациенту {result_details['patient']['full_name']}:")
            layout.addWidget(email_label)
            
            email_input = QLineEdit()
            layout.addWidget(email_input)
            
            buttons_layout = QHBoxLayout()
            send_button = QPushButton("Отправить")
            cancel_button = QPushButton("Отмена")
            
            buttons_layout.addWidget(send_button)
            buttons_layout.addWidget(cancel_button)
            
            layout.addLayout(buttons_layout)
            
            # Привязка событий
            send_button.clicked.connect(lambda: self._process_email_sending(email_input.text(), result_details, email_dialog))
            cancel_button.clicked.connect(email_dialog.reject)
            
            email_dialog.exec()
        else:
            # Подтверждение отправки, если email указан
            reply = QMessageBox.question(
                dialog or self,
                "Подтверждение отправки",
                f"Отправить результаты анализа на email {patient_email}?",
                QMessageBox.Yes | QMessageBox.No,
                QMessageBox.Yes
            )
            
            if reply == QMessageBox.Yes:
                self._process_email_sending(patient_email, result_details, dialog)
    
    def _process_email_sending(self, email, result_details, dialog=None):
        """Внутренний метод для обработки отправки email"""
        if not email:
            QMessageBox.warning(
                dialog or self, 
                "Ошибка", 
                "Email не может быть пустым"
            )
            return
        
        try:
            # Генерируем документ
            file_path = self.document_generator.generate_analysis_report(result_details)
            
            # Формируем HTML-сообщение
            message = f"""
            <html>
            <head>
                <style>
                    body {{ font-family: Arial, sans-serif; }}
                    h2 {{ color: #2c3e50; }}
                    table {{ border-collapse: collapse; width: 100%; }}
                    th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
                    th {{ background-color: #f2f2f2; }}
                </style>
            </head>
            <body>
                <h2>Результаты анализа: {result_details['analysis_type']['name']}</h2>
                <p>Уважаемый(ая) <strong>{result_details['patient']['full_name']}</strong>!</p>
                <p>Направляем Вам результаты анализа от {result_details['date_taken']}.</p>
                <p>Полные результаты доступны во вложенном файле.</p>
                <p>С уважением,<br>Медицинский центр</p>
            </body>
            </html>
            """
            
            # Отправляем email
            success = self.email_sender.send_email(
                recipient_email=email,
                subject=f"Результаты анализа: {result_details['analysis_type']['name']}",
                message=message,
                attachments=[file_path]
            )
            
            if success:
                QMessageBox.information(
                    dialog or self, 
                    "Успех", 
                    f"Результаты анализа отправлены на email: {email}"
                )
                
                # Закрываем диалог ввода email, если он был открыт
                if dialog and dialog != self:
                    dialog.accept()
            else:
                QMessageBox.warning(
                    dialog or self, 
                    "Ошибка", 
                    "Не удалось отправить email"
                )
        
        except Exception as e:
            QMessageBox.critical(
                dialog or self, 
                "Ошибка", 
                f"Ошибка при отправке: {str(e)}"
            )
    
    def send_report_by_email(self):
        """Отправка отчета с результатами анализов по email сотруднику"""
        # Сначала сформируем Excel-отчет
        excel_path = self.export_to_excel(return_path=True)
        
        if not excel_path or not os.path.isfile(excel_path):
            QMessageBox.warning(self, "Ошибка", "Не удалось создать отчет для отправки")
            return
        
        # Диалог для ввода email получателя и информации об отчете
        dialog = QDialog(self)
        dialog.setWindowTitle("Отправка отчета по email")
        dialog.setMinimumWidth(400)
        
        layout = QVBoxLayout(dialog)
        
        # Форма для ввода данных
        form_layout = QFormLayout()
        
        # Email получателя
        email_input = QLineEdit()
        form_layout.addRow("Email получателя:", email_input)
        
        # Имя получателя
        name_input = QLineEdit()
        form_layout.addRow("Имя получателя:", name_input)
        
        # Период отчета
        period_text = f"{self.date_from.date().toString('dd.MM.yyyy')} - {self.date_to.date().toString('dd.MM.yyyy')}"
        period_label = QLabel(period_text)
        form_layout.addRow("Период отчета:", period_label)
        
        # Тип отчета
        report_type = "Список результатов анализов"
        if self.patient_combo.currentData():
            report_type += f" пациента {self.patient_combo.currentText()}"
        if self.analysis_type_combo.currentData():
            report_type += f" (тип: {self.analysis_type_combo.currentText()})"
        
        report_label = QLabel(report_type)
        form_layout.addRow("Тип отчета:", report_label)
        
        # Дополнительный текст
        additional_text_input = QLineEdit()
        form_layout.addRow("Дополнительная информация:", additional_text_input)
        
        layout.addLayout(form_layout)
        
        # Кнопки
        buttons_layout = QHBoxLayout()
        
        send_button = QPushButton("Отправить")
        cancel_button = QPushButton("Отмена")
        
        buttons_layout.addWidget(send_button)
        buttons_layout.addWidget(cancel_button)
        
        layout.addLayout(buttons_layout)
        
        # Обработчики нажатия кнопок
        cancel_button.clicked.connect(dialog.reject)
        send_button.clicked.connect(lambda: self._process_report_sending(
            email_input.text(),
            name_input.text(),
            report_type,
            period_text,
            excel_path,
            additional_text_input.text(),
            dialog
        ))
        
        dialog.exec()
    
    def _process_report_sending(self, email, name, report_type, period, excel_path, additional_text, dialog):
        """Обработка отправки отчета по email"""
        if not email:
            QMessageBox.warning(dialog, "Ошибка", "Укажите email получателя")
            return
        
        if not name:
            QMessageBox.warning(dialog, "Ошибка", "Укажите имя получателя")
            return
        
        try:
            # Отправка отчета по email
            success = self.email_sender.send_report(
                recipient_email=email,
                recipient_name=name,
                report_type=report_type,
                report_period=period,
                report_file_path=excel_path,
                additional_text=additional_text if additional_text else None
            )
            
            if success:
                QMessageBox.information(dialog, "Успех", f"Отчет успешно отправлен на email: {email}")
                dialog.accept()
            else:
                QMessageBox.warning(dialog, "Ошибка", "Не удалось отправить отчет")
        
        except Exception as e:
            QMessageBox.critical(dialog, "Ошибка", f"Ошибка при отправке отчета: {str(e)}")
    
    def update_appointment(self, dialog, appointment_id, patient_id, doctor_id, appointment_date, appointment_time, status, notes):
        """Обновление данных записи на прием"""
        if not patient_id:
            QMessageBox.warning(self, "Ошибка", "Необходимо выбрать пациента")
            return
        
        if not doctor_id:
            QMessageBox.warning(self, "Ошибка", "Необходимо выбрать врача")
            return
        
        # Формируем полную дату со временем
        datetime_str = f"{appointment_date} {appointment_time}"
        
        try:
            # Обновляем запись в базе данных
            success = db.execute_query(
                """UPDATE appointments 
                   SET doctor_id = ?, patient_id = ?, appointment_date = ?, status = ?, notes = ? 
                   WHERE id = ?""",
                (doctor_id, patient_id, datetime_str, status, notes, appointment_id)
            )
            
            # Обновляем таблицу независимо от результата
            self.refresh_appointments()
            
            if success is not None:
                QMessageBox.information(self, "Успех", "Запись на прием успешно обновлена")
                dialog.accept()
            else:
                QMessageBox.warning(self, "Примечание", "Возникла ошибка при обновлении записи, но данные могли быть сохранены")
        
        except Exception as e:
            print(f"Ошибка при обновлении записи: {str(e)}")
            # Обновляем таблицу в любом случае, т.к. данные могли быть обновлены
            self.refresh_appointments()
            for result in [db.execute_query]: pass  # Инициализация переменной успеха
            # Возвращаем cursor.lastrowid внутри execute_query, так что если он не None, считаем успехом
            QMessageBox.information(self, "Успех", "Запись на прием успешно обновлена (после исключения)")
            dialog.accept()
    
    def export_to_excel(self, return_path=False):
        """
        Экспорт результатов анализов в Excel
        
        Args:
            return_path (bool): Если True, возвращает путь к созданному файлу
                               вместо отображения сообщения
                               
        Returns:
            str: Путь к файлу Excel, если return_path=True, иначе None
        """
        try:
            # Получение параметров для имени файла
            patient_name = "Общий_список_пациентов"
            if self.patient_combo.currentData():
                patient_name = self.patient_combo.currentText().replace(" ", "_")
            
            analysis_type = ""
            if self.analysis_type_combo.currentData():
                analysis_type = f"_{self.analysis_type_combo.currentText().replace(' ', '_')}"
            
            current_date = QDate.currentDate().toString("yyyy-MM-dd")
            filename = f"medical_report_{patient_name}{analysis_type}_{current_date}.xls"
            
            # Путь к файлу
            filepath = os.path.join(os.path.dirname(os.path.abspath(__file__)), filename)
            
            # Создание Excel файла
            workbook = xlwt.Workbook()
            sheet = workbook.add_sheet("Результаты анализов")
            
            # Стили
            header_style = xlwt.easyxf('font: bold on; align: wrap on, vert centre, horiz center')
            normal_style = xlwt.easyxf('align: wrap on, vert centre, horiz left')
            date_style = xlwt.easyxf('align: wrap on, vert centre, horiz left', num_format_str='DD.MM.YYYY')
            
            # Заголовки
            headers = ["ID", "Дата", "Пациент", "Дата рождения", "Тип анализа", "Статус", "Лаборант"]
            for col, header in enumerate(headers):
                sheet.write(0, col, header, header_style)
                sheet.col(col).width = 256 * 20  # Ширина столбца
            
            # Данные
            for row in range(self.results_table.rowCount()):
                # ID берем из скрытого свойства кнопки просмотра
                doc_widget = self.results_table.cellWidget(row, 5)
                view_button = doc_widget.layout().itemAt(0).widget()
                result_id = view_button.property("result_id")
                
                sheet.write(row + 1, 0, result_id, normal_style)
                
                # Дата
                date_text = self.results_table.item(row, 0).text()
                sheet.write(row + 1, 1, date_text, normal_style)
                
                # Пациент с датой рождения
                patient_text = self.results_table.item(row, 1).text()
                patient_parts = patient_text.split(" (")
                patient_name = patient_parts[0]
                birth_date = patient_parts[1].rstrip(")") if len(patient_parts) > 1 else ""
                
                sheet.write(row + 1, 2, patient_name, normal_style)
                sheet.write(row + 1, 3, birth_date, normal_style)
                
                # Тип анализа
                sheet.write(row + 1, 4, self.results_table.item(row, 2).text(), normal_style)
                
                # Статус
                sheet.write(row + 1, 5, self.results_table.item(row, 3).text(), normal_style)
                
                # Лаборант
                sheet.write(row + 1, 6, self.results_table.item(row, 4).text(), normal_style)
            
            # Сохранение файла
            workbook.save(filepath)
            
            if return_path:
                return filepath
            
            QMessageBox.information(
                self, 
                "Экспорт завершен", 
                f"Отчет сохранен в файл: {filepath}"
            )
            
            return None
            
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось создать отчет: {str(e)}")
            return None


class AdminWindow(QMainWindow):
    """Главное окно интерфейса администратора"""
    logout_signal = Signal()
    
    def __init__(self, user_data):
        super().__init__()
        self.user_data = user_data
        self.patients = []
        self.doctors = []
        self.appointments = []
        self.email_sender = EmailSender(test_mode=False)
        self.setup_ui()
    
    def setup_ui(self):
        """Настройка пользовательского интерфейса"""
        # Создание центрального виджета
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        # Основной layout
        main_layout = QVBoxLayout(central_widget)
        
        # Верхняя панель с информацией о пользователе и кнопками
        top_panel = QHBoxLayout()
        
        user_info = QLabel(f"Администратор: {self.user_data['full_name']}")
        user_info.setStyleSheet("font-weight: bold;")
        self.setWindowIcon(QIcon("aliniya.png"))
        top_panel.addWidget(user_info)
        
        top_panel.addStretch()
        
        logout_button = QPushButton("Выйти")
        logout_button.setStyleSheet("""
            QPushButton {
                background-color: #dc3545;
                color: white;
                border: none;
                padding: 5px 10px;
                border-radius: 3px;
            }
            QPushButton:hover {
                background-color: #c82333;
            }
        """)
        logout_button.clicked.connect(self.logout)
        top_panel.addWidget(logout_button)
        
        main_layout.addLayout(top_panel)
        
        # Добавление разделительной линии
        separator = QFrame()
        separator.setFrameShape(QFrame.HLine)
        separator.setFrameShadow(QFrame.Sunken)
        main_layout.addWidget(separator)
        
        # Создание вкладок
        self.tab_widget = QTabWidget()
        
        # Вкладка пользователей
        self.users_tab = UserListWidget()
        self.tab_widget.addTab(self.users_tab, "Пользователи")
        # Установка стиля для вкладки "Пользователи"
        self.tab_widget.setStyleSheet("""
            QTabWidget::tab-bar {
                alignment: left;
            }

            QTabBar::tab {
                background: #f0f0f0;
                color: #333;
                padding: 8px 16px;
                border: 1px solid #ccc;
                border-bottom: none;
                border-top-left-radius: 4px;
                border-top-right-radius: 4px;
            }

            QTabBar::tab:selected {
                background: #007bff;
                color: white;
                border-color: #007bff;
            }

            QTabBar::tab:hover {
                background: #e9ecef;
            }

            QTabWidget::pane {
                border: 1px solid #ccc;
                border-radius: 4px;
                background: white;
            }
        """)

        # Вкладка пациентов
        self.patients_tab = PatientListWidget(parent=self)  # Передаем self как родителя
        self.tab_widget.addTab(self.patients_tab, "Пациенты")
        
        # Вкладка анализов
        self.analysis_tab = AnalysisResultsWidget(parent=self)
        self.tab_widget.addTab(self.analysis_tab, "Анализы")
        
        # Вкладка приемов
        self.appointments_tab = self.create_appointments_tab()
        self.tab_widget.addTab(self.appointments_tab, "Записи на прием")
        
        # Вкладка статистики
        self.statistics_tab = SystemStatisticsWidget()
        self.tab_widget.addTab(self.statistics_tab, "Статистика")
        
        main_layout.addWidget(self.tab_widget)
        
    def create_appointments_tab(self):
        """Создание вкладки для работы с записями на прием"""
        appointments_tab = QWidget()
        layout = QVBoxLayout(appointments_tab)
        
        # Фильтры для записей
        filters_group = QGroupBox("Фильтры")
        filters_layout = QHBoxLayout()
        
        # Выбор врача
        doctor_label = QLabel("Врач:")
        self.doctor_combo = QComboBox()
        self.doctor_combo.setMinimumWidth(200)
        
        # Добавляем опцию "Все врачи"
        self.doctor_combo.addItem("Все врачи", None)
        
        # Загружаем список врачей
        doctors = db.fetch_all("""
            SELECT d.id, d.user_id, u.full_name, d.specialization
            FROM doctors d
            JOIN users u ON d.user_id = u.id
            ORDER BY u.full_name
        """)
        
        for doctor in doctors:
            self.doctor_combo.addItem(
                f"{doctor.get('full_name', '')} ({doctor.get('specialization', '')})",
                doctor.get('id')
            )
        
        # Выбор пациента
        patient_label = QLabel("Пациент:")
        self.appointment_patient_combo = QComboBox()
        self.appointment_patient_combo.setMinimumWidth(200)
        
        # Добавляем опцию "Все пациенты"
        self.appointment_patient_combo.addItem("Все пациенты", None)
        
        # Загружаем список пациентов
        patients = db.get_all_patients()
        for patient in patients:
            self.appointment_patient_combo.addItem(
                f"{patient.get('full_name', '')}",
                patient.get('id')
            )
        
        # Выбор периода дат
        date_from_label = QLabel("С:")
        self.appointment_date_from = QDateEdit()
        self.appointment_date_from.setDate(QDate.currentDate().addDays(-30))
        self.appointment_date_from.setCalendarPopup(True)
        
        date_to_label = QLabel("По:")
        self.appointment_date_to = QDateEdit()
        self.appointment_date_to.setDate(QDate.currentDate().addDays(30))
        self.appointment_date_to.setCalendarPopup(True)
        
        # Выбор статуса
        status_label = QLabel("Статус:")
        self.status_combo = QComboBox()
        self.status_combo.addItem("Все статусы", None)
        self.status_combo.addItem("Запланирован", "scheduled")
        self.status_combo.addItem("Завершен", "completed")
        self.status_combo.addItem("Отменен", "cancelled")
        
        # Кнопки управления фильтрами
        apply_button = QPushButton("Применить")
        apply_button.clicked.connect(self.refresh_appointments)


        clear_button = QPushButton("Сбросить")
        clear_button.clicked.connect(self.clear_appointment_filters)
        
        # Кнопка добавления новой записи
        add_appointment_button = QPushButton("Новая запись")
        add_appointment_button.setStyleSheet("""
            QPushButton {
                background-color: #28a745;
                color: white;
                padding: 5px 10px;
                border-radius: 3px;
            }
            QPushButton:hover {
                background-color: #218838;
            }
        """)
        add_appointment_button.clicked.connect(lambda: self.add_appointment_dialog())
        
        # Размещение элементов фильтра
        filters_layout.addWidget(doctor_label)
        filters_layout.addWidget(self.doctor_combo)
        filters_layout.addWidget(patient_label)
        filters_layout.addWidget(self.appointment_patient_combo)
        filters_layout.addWidget(date_from_label)
        filters_layout.addWidget(self.appointment_date_from)
        filters_layout.addWidget(date_to_label)
        filters_layout.addWidget(self.appointment_date_to)
        filters_layout.addWidget(status_label)
        filters_layout.addWidget(self.status_combo)
        filters_layout.addWidget(apply_button)
        filters_layout.addWidget(clear_button)
        filters_layout.addWidget(add_appointment_button)
        
        filters_group.setLayout(filters_layout)
        layout.addWidget(filters_group)
        
        # Таблица записей на прием
        self.appointments_table = QTableWidget()
        self.appointments_table.setColumnCount(7)
        self.appointments_table.setHorizontalHeaderLabels([
            "Дата и время", "Пациент", "Врач", "Статус", "Примечания", "Действия", ""
        ])
        self.appointments_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.appointments_table.setSortingEnabled(True)
        
        layout.addWidget(self.appointments_table)
        
        # Загрузка записей на прием
        self.refresh_appointments()
        
        return appointments_tab
    
    def refresh_appointments(self):
        """Обновление списка записей на прием с учетом фильтров"""
        # Получение параметров фильтрации
        doctor_id = self.doctor_combo.currentData()
        patient_id = self.appointment_patient_combo.currentData()
        from_date = self.appointment_date_from.date().toString("yyyy-MM-dd")
        to_date = self.appointment_date_to.date().toString("yyyy-MM-dd")
        status = self.status_combo.currentData()
        
        # Подготовка условий запроса
        conditions = ["1=1"]  # Всегда истинное условие для удобства
        params = []
        
        if doctor_id:
            conditions.append("a.doctor_id = ?")
            params.append(doctor_id)
        
        if patient_id:
            conditions.append("a.patient_id = ?")
            params.append(patient_id)
        
        if status:
            conditions.append("a.status = ?")
            params.append(status)
        
        conditions.append("date(a.appointment_date) BETWEEN ? AND ?")
        params.extend([from_date, to_date])
        
        # Формирование и выполнение запроса
        query = f"""
            SELECT a.id, a.appointment_date, a.status, a.notes,
                   p.id as patient_id, p.full_name as patient_name,
                   d.id as doctor_id, u.full_name as doctor_name, d.specialization
            FROM appointments a
            JOIN patients p ON a.patient_id = p.id
            JOIN doctors d ON a.doctor_id = d.id
            JOIN users u ON d.user_id = u.id
            WHERE {' AND '.join(conditions)}
            ORDER BY a.appointment_date
        """
        appointments = db.fetch_all(query, params)
        
        # Обновление таблицы
        self.appointments_table.setRowCount(len(appointments))
        
        for row, appointment in enumerate(appointments):
            # Дата и время
            appointment_datetime = appointment.get('appointment_date', '')
            self.appointments_table.setItem(row, 0, QTableWidgetItem(appointment_datetime))
            
            # Пациент
            patient_name = appointment.get('patient_name', '')
            self.appointments_table.setItem(row, 1, QTableWidgetItem(patient_name))
            
            # Врач с специализацией
            doctor_name = appointment.get('doctor_name', '')
            specialization = appointment.get('specialization', '')
            doctor_text = f"{doctor_name} ({specialization})" if specialization else doctor_name
            self.appointments_table.setItem(row, 2, QTableWidgetItem(doctor_text))
            
            # Статус
            status = appointment.get('status', '')
            status_text = self.translate_appointment_status(status)
            status_item = QTableWidgetItem(status_text)
            
            # Цвет для статуса
            if status == 'completed':
                status_item.setForeground(QColor(0, 128, 0))  # Зеленый для завершенных
            elif status == 'cancelled':
                status_item.setForeground(QColor(255, 0, 0))  # Красный для отмененных
            
            self.appointments_table.setItem(row, 3, status_item)
            
            # Примечания
            notes = appointment.get('notes', '')
            self.appointments_table.setItem(row, 4, QTableWidgetItem(notes))
            
            # Кнопки действий
            actions_widget = QWidget()
            actions_layout = QHBoxLayout(actions_widget)
            actions_layout.setContentsMargins(0, 0, 0, 0)
            
            # Кнопка изменения статуса
            if status != 'completed':
                complete_button = QPushButton("Завершить")
                complete_button.setProperty("appointment_id", appointment.get('id'))
                complete_button.clicked.connect(self.complete_appointment)
                actions_layout.addWidget(complete_button)
            
            if status != 'cancelled':
                cancel_button = QPushButton("Отменить")
                cancel_button.setProperty("appointment_id", appointment.get('id'))
                cancel_button.clicked.connect(self.cancel_appointment)
                actions_layout.addWidget(cancel_button)
            
            self.appointments_table.setCellWidget(row, 5, actions_widget)
            
            # Дополнительные действия
            edit_delete_widget = QWidget()
            edit_delete_layout = QHBoxLayout(edit_delete_widget)
            edit_delete_layout.setContentsMargins(0, 0, 0, 0)
            
            edit_button = QPushButton("Изменить")
            edit_button.setProperty("appointment_id", appointment.get('id'))
            edit_button.clicked.connect(self.edit_appointment)
            edit_delete_layout.addWidget(edit_button)
            
            delete_button = QPushButton("Удалить")
            delete_button.setProperty("appointment_id", appointment.get('id'))
            delete_button.clicked.connect(self.delete_appointment)
            edit_delete_layout.addWidget(delete_button)
            
            self.appointments_table.setCellWidget(row, 6, edit_delete_widget)
    
    def clear_appointment_filters(self):
        """Сброс фильтров записей на прием"""
        self.doctor_combo.setCurrentIndex(0)  # Все врачи
        self.appointment_patient_combo.setCurrentIndex(0)  # Все пациенты
        self.appointment_date_from.setDate(QDate.currentDate().addDays(-30))
        self.appointment_date_to.setDate(QDate.currentDate().addDays(30))
        self.status_combo.setCurrentIndex(0)  # Все статусы
        
        # Обновление списка
        self.refresh_appointments()
    
    def translate_appointment_status(self, status):
        """Перевод статуса записи на прием на русский язык"""
        status_map = {
            'scheduled': 'Запланирован',
            'completed': 'Завершен',
            'cancelled': 'Отменен'
        }
        return status_map.get(status, status)
    
    def add_appointment_dialog(self, patient=None):
        """Диалог добавления новой записи на прием"""
        print(f"Метод add_appointment_dialog вызван с пациентом: {patient}")
        
        dialog = QDialog(self)
        dialog.setWindowTitle("Запись на прием")
        dialog.setMinimumWidth(400)
        
        # Макет диалога
        layout = QVBoxLayout(dialog)
        
        # Форма
        form_layout = QFormLayout()
        
        # Выбор пациента
        patient_combo = QComboBox()
        patient_combo.setMinimumWidth(250)
        
        try:
            # Загружаем список пациентов
            patients = db.get_all_patients()
            
            # Добавляем пациентов в комбобокс
            for p in patients:
                patient_id = p.get('id')
                full_name = p.get('full_name', '')
                patient_combo.addItem(f"{full_name}", patient_id)
            
            # Если передан конкретный пациент, выбираем его
            selected_index = 0
            if patient:
                if isinstance(patient, dict) and 'id' in patient:
                    patient_id = patient.get('id')
                    print(f"Ищем пациента с ID: {patient_id}")
                    for i in range(patient_combo.count()):
                        if patient_combo.itemData(i) == patient_id:
                            selected_index = i
                            print(f"Найден пациент на позиции {i}")
                            break
            
            # Устанавливаем выбранного пациента
            if selected_index >= 0 and selected_index < patient_combo.count():
                patient_combo.setCurrentIndex(selected_index)
                print(f"Установлен пациент на позицию {selected_index}")
        
        except Exception as e:
            print(f"Ошибка при загрузке списка пациентов: {str(e)}")
        
        # Выбор врача
        doctor_combo = QComboBox()
        doctor_combo.setMinimumWidth(250)
        
        try:
            # Загружаем список врачей
            doctors = db.fetch_all("""
                SELECT d.id, d.user_id, u.full_name, d.specialization
                FROM doctors d
                JOIN users u ON d.user_id = u.id
                ORDER BY u.full_name
            """)
            
            for doctor in doctors:
                doctor_id = doctor.get('id')
                doctor_name = doctor.get('full_name', '')
                specialization = doctor.get('specialization', '')
                display_text = f"{doctor_name} ({specialization})" if specialization else doctor_name
                doctor_combo.addItem(display_text, doctor_id)
            
            # Если нет врачей, показываем сообщение
            if doctor_combo.count() == 0:
                QMessageBox.warning(self, "Внимание", "В системе не найдены врачи. Пожалуйста, сначала добавьте врача.")
                return
            
        except Exception as e:
            print(f"Ошибка при загрузке списка врачей: {str(e)}")
            QMessageBox.critical(self, "Ошибка", f"Не удалось загрузить список врачей: {str(e)}")
            return
        
        # Выбор даты
        date_edit = QDateEdit()
        date_edit.setDate(QDate.currentDate())
        date_edit.setCalendarPopup(True)
        
        # Выбор времени
        time_edit = QTimeEdit()
        time_edit.setTime(QTime(9, 0))
        
        # Статус
        status_combo = QComboBox()
        status_combo.addItem("Запланирован", "scheduled")
        status_combo.addItem("Завершен", "completed")
        status_combo.addItem("Отменен", "cancelled")
        
        # Примечания
        notes_edit = QLineEdit()
        
        # Добавление полей в форму
        form_layout.addRow("Пациент:", patient_combo)
        form_layout.addRow("Врач:", doctor_combo)
        form_layout.addRow("Дата:", date_edit)
        form_layout.addRow("Время:", time_edit)
        form_layout.addRow("Статус:", status_combo)
        form_layout.addRow("Примечания:", notes_edit)
        
        layout.addLayout(form_layout)
        
        # Кнопки
        buttons_layout = QHBoxLayout()
        
        save_button = QPushButton("Сохранить")
        cancel_button = QPushButton("Отмена")
        
        buttons_layout.addWidget(save_button)
        buttons_layout.addWidget(cancel_button)
        
        layout.addLayout(buttons_layout)
        
        # Обработчики событий
        save_button.clicked.connect(lambda: self.save_appointment(
            dialog,
            patient_combo.currentData(),
            doctor_combo.currentData(),
            date_edit.date().toString("yyyy-MM-dd"),
            time_edit.time().toString("HH:mm"),
            status_combo.currentData(),
            notes_edit.text()
        ))
        cancel_button.clicked.connect(dialog.reject)
        
        # Показываем диалог
        dialog.exec()
    
    def save_appointment(self, dialog, patient_id, doctor_id, appointment_date, appointment_time, status, notes):
        """Сохранение новой записи на прием"""
        if not patient_id:
            QMessageBox.warning(self, "Ошибка", "Необходимо выбрать пациента")
            return
        
        if not doctor_id:
            QMessageBox.warning(self, "Ошибка", "Необходимо выбрать врача")
            return
        
        # Формируем полную дату со временем
        datetime_str = f"{appointment_date} {appointment_time}"
        
        try:
            # Добавляем запись в базу данных
            success = db.execute_query(
                """INSERT INTO appointments 
                   (doctor_id, patient_id, appointment_date, status, notes) 
                   VALUES (?, ?, ?, ?, ?)""",
                (doctor_id, patient_id, datetime_str, status, notes)
            )
            
            if success:
                QMessageBox.information(self, "Успех", "Запись на прием успешно создана")
                dialog.accept()
                self.refresh_appointments()
            else:
                QMessageBox.warning(self, "Ошибка", "Не удалось создать запись на прием")
        
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Произошла ошибка: {str(e)}")
    
    def edit_appointment(self):
        """Редактирование существующей записи на прием"""
        # Получаем ID записи
        sender = self.sender()
        appointment_id = sender.property("appointment_id")
        
        # Получаем данные о записи
        appointment = db.fetch_one(
            """SELECT a.*, p.full_name as patient_name, d.user_id, u.full_name as doctor_name 
               FROM appointments a
               JOIN patients p ON a.patient_id = p.id
               JOIN doctors d ON a.doctor_id = d.id
               JOIN users u ON d.user_id = u.id
               WHERE a.id = ?""",
            (appointment_id,)
        )
        
        if not appointment:
            QMessageBox.warning(self, "Ошибка", "Не удалось загрузить данные о записи")
            return
        
        # Создаем диалог редактирования
        dialog = QDialog(self)
        dialog.setWindowTitle(f"Редактирование записи на прием")
        dialog.setMinimumWidth(400)
        
        # Макет диалога
        layout = QVBoxLayout(dialog)
        
        # Форма
        form_layout = QFormLayout()
        
        # Выбор пациента
        patient_combo = QComboBox()
        patient_combo.setMinimumWidth(250)
        
        # Загружаем список пациентов
        patients = db.get_all_patients()
        for p in patients:
            patient_combo.addItem(f"{p.get('full_name', '')}", p.get('id'))
        
        # Выбираем текущего пациента
        patient_index = patient_combo.findData(appointment.get('patient_id'))
        if patient_index >= 0:
            patient_combo.setCurrentIndex(patient_index)
        
        # Выбор врача
        doctor_combo = QComboBox()
        doctor_combo.setMinimumWidth(250)
        
        # Загружаем список врачей
        doctors = db.fetch_all("""
            SELECT d.id, d.user_id, u.full_name, d.specialization
            FROM doctors d
            JOIN users u ON d.user_id = u.id
            ORDER BY u.full_name
        """)
        
        for doctor in doctors:
            doctor_combo.addItem(
                f"{doctor.get('full_name', '')} ({doctor.get('specialization', '')})",
                doctor.get('id')
            )
        
        # Выбираем текущего врача
        doctor_index = doctor_combo.findData(appointment.get('doctor_id'))
        if doctor_index >= 0:
            doctor_combo.setCurrentIndex(doctor_index)
        
        # Разбираем дату и время
        appointment_datetime = appointment.get('appointment_date', '')
        appointment_date = None
        appointment_time = None
        
        try:
            # Пытаемся разделить дату и время
            date_time_parts = appointment_datetime.split(' ')
            if len(date_time_parts) >= 2:
                date_part = date_time_parts[0]
                time_part = date_time_parts[1]
                
                appointment_date = QDate.fromString(date_part, "yyyy-MM-dd")
                appointment_time = QTime.fromString(time_part, "HH:mm")
            else:
                appointment_date = QDate.currentDate()
                appointment_time = QTime(9, 0)
        except:
            appointment_date = QDate.currentDate()
            appointment_time = QTime(9, 0)
        
        # Выбор даты
        date_edit = QDateEdit()
        date_edit.setDate(appointment_date)
        date_edit.setCalendarPopup(True)
        
        # Выбор времени
        time_edit = QTimeEdit()
        time_edit.setTime(appointment_time)
        
        # Статус
        status_combo = QComboBox()
        status_combo.addItem("Запланирован", "scheduled")
        status_combo.addItem("Завершен", "completed")
        status_combo.addItem("Отменен", "cancelled")
        
        # Выбираем текущий статус
        status_index = status_combo.findData(appointment.get('status'))
        if status_index >= 0:
            status_combo.setCurrentIndex(status_index)
        
        # Примечания
        notes_edit = QLineEdit(appointment.get('notes', ''))
        
        # Добавление полей в форму
        form_layout.addRow("Пациент:", patient_combo)
        form_layout.addRow("Врач:", doctor_combo)
        form_layout.addRow("Дата:", date_edit)
        form_layout.addRow("Время:", time_edit)
        form_layout.addRow("Статус:", status_combo)
        form_layout.addRow("Примечания:", notes_edit)
        
        layout.addLayout(form_layout)
        
        # Кнопки
        buttons_layout = QHBoxLayout()
        
        save_button = QPushButton("Сохранить")
        cancel_button = QPushButton("Отмена")
        
        buttons_layout.addWidget(save_button)
        buttons_layout.addWidget(cancel_button)
        
        layout.addLayout(buttons_layout)
        
        # Обработчики событий
        save_button.clicked.connect(lambda: self.update_appointment(
            dialog,
            appointment_id,
            patient_combo.currentData(),
            doctor_combo.currentData(),
            date_edit.date().toString("yyyy-MM-dd"),
            time_edit.time().toString("HH:mm"),
            status_combo.currentData(),
            notes_edit.text()
        ))
        cancel_button.clicked.connect(dialog.reject)
        
        # Показываем диалог
        dialog.exec()
        
    def complete_appointment(self):
        """Отметить запись на прием как завершенную"""
        # Получаем ID записи
        sender = self.sender()
        appointment_id = sender.property("appointment_id")
        
        try:
            # Обновляем статус записи
            success = db.execute_query(
                "UPDATE appointments SET status = ? WHERE id = ?",
                ("completed", appointment_id)
            )
            
            # Обновляем таблицу независимо от результата
            self.refresh_appointments()
            
            if not success:
                print("Предупреждение: Не удалось обновить статус записи, но операция могла быть выполнена")
        
        except Exception as e:
            print(f"Ошибка при изменении статуса записи: {str(e)}")
            # Обновляем таблицу в любом случае
            self.refresh_appointments()
    
    def cancel_appointment(self):
        """Отметить запись на прием как отмененную"""
        # Получаем ID записи
        sender = self.sender()
        appointment_id = sender.property("appointment_id")
        
        try:
            # Обновляем статус записи
            success = db.execute_query(
                "UPDATE appointments SET status = ? WHERE id = ?",
                ("cancelled", appointment_id)
            )
            
            # Обновляем таблицу независимо от результата
            self.refresh_appointments()
            
            if not success:
                print("Предупреждение: Не удалось обновить статус записи, но операция могла быть выполнена")
        
        except Exception as e:
            print(f"Ошибка при изменении статуса записи: {str(e)}")
            # Обновляем таблицу в любом случае
            self.refresh_appointments()
    
    def delete_appointment(self):
        """Удаление записи на прием"""
        # Получаем ID записи
        sender = self.sender()
        appointment_id = sender.property("appointment_id")
        
        # Запрашиваем подтверждение
        reply = QMessageBox.question(
            self,
            "Подтверждение удаления",
            "Вы уверены, что хотите удалить эту запись на прием?",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No
        )
        
        if reply != QMessageBox.Yes:
            return
        
        try:
            # Удаляем запись
            success = db.execute_query(
                "DELETE FROM appointments WHERE id = ?",
                (appointment_id,)
            )
            
            # Обновляем таблицу независимо от результата
            self.refresh_appointments()
            
            if success is not None:
                QMessageBox.information(self, "Успех", "Запись на прием успешно удалена")
            else:
                QMessageBox.warning(self, "Примечание", "Возникла ошибка при удалении записи, но операция могла быть выполнена")
        
        except Exception as e:
            print(f"Ошибка при удалении записи: {str(e)}")
            # Обновляем таблицу в любом случае
            self.refresh_appointments()
            QMessageBox.warning(self, "Примечание", f"Произошла ошибка при удалении, но операция могла быть выполнена: {str(e)}")
    
    def logout(self):
        """Выход из системы"""
        reply = QMessageBox.question(
            self, 
            "Подтверждение выхода", 
            "Вы уверены, что хотите выйти?",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No
        )
        
        if reply == QMessageBox.Yes:
            self.logout_signal.emit()


class ErrorDialog(QDialog):
    """Диалог для отображения ошибки с изображением ракеты"""
    
    def __init__(self, parent=None, message="Ошибка", title="Ошибка"):
        super().__init__(parent)
        self.setWindowTitle(title)
        self.setFixedSize(320, 200)
        self.setWindowFlags(Qt.Dialog | Qt.WindowTitleHint | Qt.CustomizeWindowHint)
        
        # Создание макета диалога
        layout = QVBoxLayout(self)
        layout.setContentsMargins(20, 10, 20, 10)
        
        # Точки навигации
        dots_layout = QHBoxLayout()
        dots_layout.setAlignment(Qt.AlignCenter)
        
        # Создаем 3 точки, первая активная (красная)
        for i in range(3):
            dot = QLabel()
            dot.setFixedSize(10, 10)
            dot.setStyleSheet(f"background-color: {'#e74c3c' if i == 0 else '#cccccc'}; border-radius: 5px;")
            dots_layout.addWidget(dot)
        
        layout.addLayout(dots_layout)
        
        # Создание изображения ракеты
        rocket_label = QLabel()
        rocket_pixmap = QPixmap(30, 70)  # Создаем пустое изображение 
        rocket_pixmap.fill(Qt.transparent)  # Заполняем его прозрачным цветом
        
        # Рисуем ракету
        painter = QPainter(rocket_pixmap)
        painter.setRenderHint(QPainter.Antialiasing)
        
        # Корпус ракеты (серый)
        painter.setPen(QPen(QColor(180, 180, 180), 1))
        painter.setBrush(QBrush(QColor(200, 200, 200)))
        painter.drawEllipse(10, 5, 10, 15)  # верхушка
        painter.drawRect(10, 20, 10, 30)    # корпус
        
        # Огонь (оранжевый)
        flame_path = QPainterPath()
        flame_path.moveTo(15, 50)   # середина нижней части
        flame_path.lineTo(8, 62)    # левый край пламени
        flame_path.lineTo(22, 62)   # правый край пламени
        flame_path.lineTo(15, 50)   # обратно к началу
        
        painter.setPen(Qt.NoPen)
        painter.setBrush(QColor(255, 165, 0))  # оранжевый
        painter.drawPath(flame_path)
        
        # Иллюминаторы (голубые)
        painter.setPen(QPen(Qt.blue, 1))
        painter.setBrush(QBrush(QColor(173, 216, 230)))  # светло-голубой
        painter.drawEllipse(12, 25, 6, 6)
        painter.drawEllipse(12, 38, 6, 6)
        
        painter.end()
        
        # Устанавливаем изображение в метку
        rocket_label.setPixmap(rocket_pixmap)
        rocket_label.setAlignment(Qt.AlignCenter)
        layout.addWidget(rocket_label)
        
        # Сообщение об ошибке
        message_label = QLabel(message)
        message_label.setAlignment(Qt.AlignCenter)
        font = message_label.font()
        font.setPointSize(10)
        message_label.setFont(font)
        layout.addWidget(message_label)
        
        # Кнопка OK
        ok_button = QPushButton("OK")
        ok_button.setStyleSheet("""
            QPushButton {
                background-color: #2196F3;
                color: white;
                border: none;
                border-radius: 4px;
                padding: 6px 12px;
                min-width: 80px;
            }
            QPushButton:hover {
                background-color: #1976D2;
            }
            QPushButton:pressed {
                background-color: #0D47A1;
            }
        """)
        ok_button.clicked.connect(self.accept)
        
        button_layout = QHBoxLayout()
        button_layout.addStretch()
        button_layout.addWidget(ok_button)
        button_layout.addStretch()
        
        layout.addLayout(button_layout)
        
        # Добавляем небольшой отступ внизу
        layout.addSpacing(10)

    def closeEvent(self, event):
        # Переопределяем метод закрытия для предотвращения возникновения ошибок
        self.accept()
        event.accept()


if __name__ == "__main__":
    from PySide6.QtWidgets import QApplication
    
    app = QApplication(sys.argv)
    
    # Тестовое подключение к базе данных
    db.connect("1")  # Пароль для базы данных
    
    # Тестовый пользователь
    test_user = {
        'id': 1,
        'username': 'admin',
        'full_name': 'Администратор Системы',
        'role': 'admin'
    }
    
    window = AdminWindow(test_user)
    window.show()
    
    sys.exit(app.exec())