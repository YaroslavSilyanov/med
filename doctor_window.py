from PyQt6.QtCore import QTime, QDateTime
from PySide6.QtWidgets import (QMainWindow, QWidget, QLabel, QComboBox, QPushButton,
                               QVBoxLayout, QHBoxLayout, QMessageBox, QFormLayout,
                               QTableWidget, QTableWidgetItem, QLineEdit, QDialog,
                               QTabWidget, QCalendarWidget, QDateEdit, QGroupBox,
                               QScrollArea, QFrame, QHeaderView, QTextEdit)
from PySide6.QtCore import Qt, Signal, QDate
from PySide6.QtGui import QFont, QIcon, QColor
import sys
import json
from datetime import datetime
import docx
from docx.shared import Pt
import os

from database_connection import db

# Единый стиль для всего приложения
GLOBAL_STYLESHEET = """
    /* Общие стили для всех виджетов */
    QWidget {
        font-family: 'Arial', sans-serif;
        color: #2c3e50;
    }

    /* Стили для QGroupBox */
    QGroupBox {
        font-weight: bold;
        font-size: 14px;
        border: 1px solid #dfe6e9;
        border-radius: 5px;
        margin-top: 20px;
        background-color: #f8f9fa;
    }

    QGroupBox::title {
        subcontrol-origin: margin;
        subcontrol-position: top left;
        padding: 5px 10px;
        background-color: #f8f9fa;
    }

    /* Стили для меток */
    QLabel {
        font-size: 12px;
        color: #2c3e50;
    }

    /* Стили для выпадающих списков */
    QComboBox {
        border: 1px solid #ced4da;
        border-radius: 4px;
        padding: 5px 10px;
        background-color: white;
        font-size: 12px;
    }

    QComboBox::drop-down {
        border-left: 1px solid #ced4da;
        padding: 0 5px;
    }

    QComboBox::down-arrow {
        width: 10px;
        height: 10px;
    }

    /* Стили для полей ввода и QDateEdit */
    QLineEdit, QDateEdit, QTextEdit {
        border: 1px solid #ced4da;
        border-radius: 4px;
        padding: 5px 10px;
        font-size: 12px;
        background-color: white;
    }

    /* Стили для таблицы */
    QTableWidget {
        border: 1px solid #ced4da;
        border-radius: 4px;
        background-color: white;
        font-size: 14px;
    }

    QTableWidget::item {
        padding: 5px;
    }

    QHeaderView::section {
        background-color: #e9ecef;
        padding: 5px;
        border: 1px solid #ced4da;
        font-weight: bold;
    }

    /* Стили для вкладок */
    QTabWidget::pane {
        border: 1px solid #ced4da;
        border-radius: 4px;
        background-color: #f8f9fa;
    }

    QTabBar::tab {
        background-color: #e9ecef;
        border: 1px solid #ced4da;
        padding: 8px 16px;
        border-top-left-radius: 4px;
        border-top-right-radius: 4px;
    }

    QTabBar::tab:selected {
        background-color: #f8f9fa;
        font-weight: bold;
    }

    /* Общие стили для кнопок */
    QPushButton {
        border-radius: 5px;
        padding: 8px 16px;
        font-size: 12px;
        font-weight: bold;
    }

    /* Основная кнопка (синяя) */
    QPushButton#primary {
        background-color: #007bff;
        color: white;
        border: none;
    }

    QPushButton#primary:hover {
        background-color: #0069d9;
    }

    /* Вторичная кнопка (серая) */
    QPushButton#secondary {
        background-color: #6c757d;
        color: white;
        border: none;
    }

    QPushButton#secondary:hover {
        background-color: #5a6268;
    }

    /* Кнопка успеха (зеленая) */
    QPushButton#success {
        background-color: #28a745;
        color: white;
        border: none;
    }

    QPushButton#success:hover {
        background-color: #218838;
    }

    /* Кнопка опасности (красная) */
    QPushButton#danger {
        background-color: #dc3545;
        color: white;
        border: none;
    }

    QPushButton#danger:hover {
        background-color: #c82333;
    }

    /* Кнопка информации (голубая) */
    QPushButton#info {
        background-color: #17a2b8;
        color: white;
        border: none;
    }

    QPushButton#info:hover {
        background-color: #138496;
    }

    /* Стили для QFrame */
    QFrame#infoFrame {
        background-color: #f8f9fa;
        border: 1px solid #dfe6e9;
        border-radius: 5px;
        padding: 10px;
    }
"""


class AppointmentDetailsDialog(QDialog):
    """Диалоговое окно с деталями приема"""

    def __init__(self, appointment_data, parent=None):
        super().__init__(parent)
        self.appointment_data = appointment_data

        self.setWindowTitle("Детали приема")
        self.setMinimumWidth(400)
        self.setup_ui()

    def setup_ui(self):
        """Настройка пользовательского интерфейса"""
        layout = QVBoxLayout()

        # Информация о пациенте
        patient_group = QGroupBox("Информация о пациенте")
        patient_layout = QFormLayout()

        patient_name = QLabel(self.appointment_data.get('patient_name', ''))
        patient_name.setFont(QFont("Arial", 12, QFont.Bold))
        patient_layout.addRow("Пациент:", patient_name)

        # Получаем дополнительную информацию о пациенте
        patient_id = self.appointment_data.get('patient_id')
        if patient_id:
            patient_info = db.get_patient(patient_id)
            if patient_info:
                patient_layout.addRow("Дата рождения:", QLabel(str(patient_info.get('birth_date', ''))))
                patient_layout.addRow("Телефон:", QLabel(patient_info.get('phone', '')))
                patient_layout.addRow("Email:", QLabel(patient_info.get('email', '')))

        patient_group.setLayout(patient_layout)
        layout.addWidget(patient_group)

        # Информация о приеме
        appointment_group = QGroupBox("Информация о приеме")
        appointment_layout = QFormLayout()

        date_time = QLabel(self.appointment_data.get('appointment_date', ''))
        appointment_layout.addRow("Дата и время:", date_time)

        status_label = QLabel(self.get_status_text(self.appointment_data.get('status', '')))
        appointment_layout.addRow("Статус:", status_label)

        notes_label = QLabel(self.appointment_data.get('notes', ''))
        notes_label.setWordWrap(True)
        appointment_layout.addRow("Примечания:", notes_label)

        appointment_group.setLayout(appointment_layout)
        layout.addWidget(appointment_group)

        # Получаем анализы пациента, если есть
        patient_id = self.appointment_data.get('patient_id')
        if patient_id:
            analysis_results = db.get_patient_analysis_results(patient_id)
            if analysis_results:
                analysis_group = QGroupBox("Анализы пациента")
                analysis_layout = QVBoxLayout()

                analysis_table = QTableWidget()
                analysis_table.setColumnCount(3)
                analysis_table.setHorizontalHeaderLabels(["Тип анализа", "Дата", "Статус"])
                analysis_table.setRowCount(len(analysis_results))

                for row, result in enumerate(analysis_results):
                    analysis_table.setItem(row, 0, QTableWidgetItem(result.get('analysis_name', '')))
                    date_str = result.get('result_date', '')
                    if not isinstance(date_str, str):
                        date_str = date_str.strftime('%d.%m.%Y %H:%M')
                    analysis_table.setItem(row, 1, QTableWidgetItem(date_str))
                    status_item = QTableWidgetItem(self.get_analysis_status(result.get('status', '')))
                    analysis_table.setItem(row, 2, status_item)

                analysis_table.resizeColumnsToContents()
                analysis_layout.addWidget(analysis_table)
                analysis_group.setLayout(analysis_layout)
                layout.addWidget(analysis_group)

        # Кнопки действий
        buttons_layout = QHBoxLayout()

        change_status_button = QPushButton("Изменить статус")
        change_status_button.setObjectName("success")
        change_status_button.clicked.connect(self.change_appointment_status)
        buttons_layout.addWidget(change_status_button)

        close_button = QPushButton("Закрыть")
        close_button.setObjectName("secondary")
        close_button.clicked.connect(self.accept)
        buttons_layout.addWidget(close_button)

        layout.addLayout(buttons_layout)

        self.setLayout(layout)

    def get_status_text(self, status):
        """Возвращает текстовое представление статуса приема"""
        status_map = {
            'scheduled': 'Запланирован',
            'completed': 'Завершен',
            'cancelled': 'Отменен'
        }
        return status_map.get(status, status)

    def get_analysis_status(self, status):
        """Возвращает текстовое представление статуса анализа"""
        status_map = {
            'pending': 'В обработке',
            'completed': 'Выполнен',
            'sent': 'Отправлен'
        }
        return status_map.get(status, status)

    def change_appointment_status(self):
        """Изменение статуса приема"""
        appointment_id = self.appointment_data.get('id')
        current_status = self.appointment_data.get('status')

        # Варианты статусов
        if current_status == 'scheduled':
            new_status = 'completed'
            status_text = 'Завершен'
        elif current_status == 'completed':
            new_status = 'scheduled'
            status_text = 'Запланирован'
        else:
            new_status = 'scheduled'
            status_text = 'Запланирован'

        reply = QMessageBox.question(
            self,
            "Изменение статуса",
            f"Изменить статус приема на '{status_text}'?",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No
        )

        if reply == QMessageBox.Yes:
            if db.update_appointment_status(appointment_id, new_status):
                QMessageBox.information(self, "Успех", "Статус приема успешно изменен")
                self.appointment_data['status'] = new_status
                self.accept()  # Закрываем окно
            else:
                QMessageBox.critical(self, "Ошибка", "Не удалось изменить статус приема")


class AnalysisDetailsDialog(QDialog):
    """Диалоговое окно с деталями анализа"""

    def __init__(self, analysis_data, parent=None):
        super().__init__(parent)
        self.analysis_data = analysis_data

        self.setWindowTitle(f"Детали анализа: {analysis_data.get('analysis_name', '')}")
        self.setMinimumWidth(500)
        self.setup_ui()

    def setup_ui(self):
        """Настройка пользовательского интерфейса"""
        layout = QVBoxLayout()

        # Информация о пациенте и анализе
        info_frame = QFrame()
        info_frame.setObjectName("infoFrame")
        info_layout = QFormLayout(info_frame)

        patient_name = QLabel(self.analysis_data.get('patient_name', ''))
        patient_name.setFont(QFont("Arial", 12, QFont.Bold))
        info_layout.addRow("Пациент:", patient_name)

        info_layout.addRow("Тип анализа:", QLabel(self.analysis_data.get('analysis_name', '')))

        date_str = self.analysis_data.get('result_date', '')
        if not isinstance(date_str, str):
            date_str = date_str.strftime('%d.%m.%Y %H:%M')
        info_layout.addRow("Дата:", QLabel(date_str))

        status_text = {
            'pending': 'В обработке',
            'completed': 'Выполнен',
            'sent': 'Отправлен'
        }.get(self.analysis_data.get('status', ''), self.analysis_data.get('status', ''))

        info_layout.addRow("Статус:", QLabel(status_text))

        layout.addWidget(info_frame)

        # Результаты анализа
        results_group = QGroupBox("Результаты анализа")
        results_layout = QVBoxLayout()

        # Парсинг данных результата
        result_data = self.analysis_data.get('result_data', '{}')
        if isinstance(result_data, str):
            try:
                result_data = json.loads(result_data)
            except json.JSONDecodeError:
                result_data = {"Результат": result_data}

        # Отображение результатов в таблице
        if isinstance(result_data, dict) and result_data:
            results_table = QTableWidget()
            results_table.setColumnCount(3)
            results_table.setHorizontalHeaderLabels(["Параметр", "Значение", "Норма"])
            results_table.setRowCount(len(result_data))

            for row, (param, value) in enumerate(result_data.items()):
                param_item = QTableWidgetItem(param)
                value_item = QTableWidgetItem(str(value))

                # Нормальные значения по справочнику
                normal_value = self.get_normal_value(param)
                normal_item = QTableWidgetItem(normal_value)

                # Функция проверки значения
                def is_value_normal(val, norm):
                    import re
                    try:
                        if norm in ['Отсутствует', 'Отсутствуют']:
                            return str(val).lower() in ['отсутствует', 'отсутствуют', 'нет', '0']
                        match = re.match(r"([\d.]+)-([\d.]+)", norm.replace(',', '.'))
                        if match:
                            min_val, max_val = map(float, match.groups())
                            return min_val <= float(val) <= max_val
                    except:
                        return False
                    return False

                # Если не в норме — выделяем красным
                if not is_value_normal(value, normal_value):
                    value_item.setForeground(QColor("red"))

                results_table.setItem(row, 0, param_item)
                results_table.setItem(row, 1, value_item)
                results_table.setItem(row, 2, normal_item)

            results_table.resizeColumnsToContents()
            results_layout.addWidget(results_table)
        else:
            results_layout.addWidget(QLabel("Нет данных о результатах анализа"))

        results_group.setLayout(results_layout)
        layout.addWidget(results_group)

        # Кнопка закрытия
        close_button = QPushButton("Закрыть")
        close_button.setObjectName("secondary")
        close_button.clicked.connect(self.accept)
        layout.addWidget(close_button)

        self.setLayout(layout)

    def get_normal_value(self, parameter):
        """Получение нормальных значений для параметра анализа"""
        normal_values = {
            'Гемоглобин': '120-160 г/л',
            'Эритроциты': '3.8-5.5 млн/мкл',
            'Лейкоциты': '4.0-9.0 тыс/мкл',
            'Тромбоциты': '180-320 тыс/мкл',
            'СОЭ': '2-15 мм/ч',
            'Глюкоза': '3.9-6.1 ммоль/л',
            'Холестерин': '3.0-5.2 ммоль/л',
            'Билирубин': '3.4-17.1 мкмоль/л',
            'АЛТ': '5-40 ед/л',
            'АСТ': '5-40 ед/л',
            'Креатинин': '53-106 мкмоль/л',
            'Мочевина': '2.5-8.3 ммоль/л',
            'pH': '5.0-7.0',
            'Белок': 'Отсутствует',
            'Кетоновые тела': 'Отсутствует',
            'Цвет': 'Светло-желтый',
            'Прозрачность': 'Прозрачная'
        }
        return normal_values.get(parameter, 'Не определено')


class PrescriptionDialog(QDialog):
    """Диалоговое окно для создания рецепта"""
    def __init__(self, patients, doctor_info, parent=None):
        super().__init__(parent)
        self.patients = patients
        self.doctor_info = doctor_info
        self.medications = db.get_all_medications()  # Получаем список лекарств
        self.medication_rows = []  # Список для хранения строк с лекарствами
        self.setWindowTitle("Выписать рецепт")
        self.setMinimumWidth(600)
        self.setup_ui()

    def setup_ui(self):
        """Настройка пользовательского интерфейса"""
        layout = QVBoxLayout()

        # Форма для основных данных рецепта
        form_group = QGroupBox("Данные рецепта")
        form_layout = QFormLayout()

        # Выбор пациента
        self.patient_combo = QComboBox()
        for patient in self.patients:
            self.patient_combo.addItem(patient['full_name'], patient['id'])
        form_layout.addRow("Пациент:", self.patient_combo)

        # Дата выдачи
        self.issue_date = QDateEdit()
        self.issue_date.setDate(QDate.currentDate())
        self.issue_date.setCalendarPopup(True)
        form_layout.addRow("Дата выдачи:", self.issue_date)

        form_group.setLayout(form_layout)
        layout.addWidget(form_group)

        # Группа для выбора лекарств
        medications_group = QGroupBox("Лекарства")
        self.medications_layout = QVBoxLayout()

        # Кнопка для добавления лекарства
        add_med_button = QPushButton("Добавить лекарство")
        add_med_button.setObjectName("primary")
        add_med_button.clicked.connect(self.add_medication_row)
        self.medications_layout.addWidget(add_med_button)

        medications_group.setLayout(self.medications_layout)
        layout.addWidget(medications_group)

        # Кнопки действий
        buttons_layout = QHBoxLayout()
        save_button = QPushButton("Сохранить")
        save_button.setObjectName("success")
        save_button.clicked.connect(self.save_prescription)
        buttons_layout.addWidget(save_button)

        cancel_button = QPushButton("Отмена")
        cancel_button.setObjectName("secondary")
        cancel_button.clicked.connect(self.reject)
        buttons_layout.addWidget(cancel_button)

        layout.addLayout(buttons_layout)
        self.setLayout(layout)

        # Добавляем одну строку для лекарства по умолчанию
        self.add_medication_row()

    def add_medication_row(self):
        """Добавление строки для выбора лекарства, дозировки и инструкций"""
        medication_row = QWidget()
        row_layout = QHBoxLayout()

        # Выбор лекарства
        medication_combo = QComboBox()
        for med in self.medications:
            medication_combo.addItem(med['name'], med['id'])
        row_layout.addWidget(medication_combo)

        # Дозировка
        dosage_input = QLineEdit()
        dosage_input.setPlaceholderText("Введите дозировку")
        row_layout.addWidget(dosage_input)

        # Инструкции
        instructions_input = QLineEdit()
        instructions_input.setPlaceholderText("Введите инструкции")
        row_layout.addWidget(instructions_input)

        # Кнопка удаления строки
        remove_button = QPushButton("Удалить")
        remove_button.setObjectName("danger")
        remove_button.clicked.connect(lambda: self.remove_medication_row(medication_row))
        row_layout.addWidget(remove_button)

        medication_row.setLayout(row_layout)
        self.medications_layout.insertWidget(self.medications_layout.count() - 1, medication_row)
        self.medication_rows.append({
            'widget': medication_row,
            'combo': medication_combo,
            'dosage': dosage_input,
            'instructions': instructions_input
        })

    def remove_medication_row(self, widget):
        """Удаление строки с лекарством"""
        for row in self.medication_rows:
            if row['widget'] == widget:
                self.medications_layout.removeWidget(widget)
                widget.deleteLater()
                self.medication_rows.remove(row)
                break

    def save_prescription(self):
        """Сохранение рецепта в базе данных"""
        patient_id = self.patient_combo.currentData()
        issue_date = self.issue_date.date().toString("yyyy-MM-dd")

        if not patient_id or not issue_date:
            QMessageBox.warning(self, "Предупреждение", "Пожалуйста, выберите пациента и дату выдачи.")
            return

        if not self.medication_rows:
            QMessageBox.warning(self, "Предупреждение", "Добавьте хотя бы одно лекарство.")
            return

        for row in self.medication_rows:
            if not row['dosage'].text().strip():
                QMessageBox.warning(self, "Предупреждение", "Укажите дозировку для всех лекарств.")
                return

        try:
            # Создаем запись рецепта
            prescription_id = db.add_prescription(
                self.doctor_info['id'],
                patient_id,
                issue_date
            )

            # Добавляем лекарства к рецепту
            for row in self.medication_rows:
                medication_id = row['combo'].currentData()
                dosage = row['dosage'].text().strip()
                instructions = row['instructions'].text().strip() or None
                db.add_prescription_medication(prescription_id, medication_id, dosage, instructions)

            QMessageBox.information(self, "Успех", "Рецепт успешно сохранен")
            self.accept()
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось сохранить рецепт: {str(e)}")

class ScheduleAppointmentDialog(QDialog):
    """Диалоговое окно для планирования приема"""

    def __init__(self, patients, doctor_info, parent=None):
        super().__init__(parent)
        self.patients = patients
        self.doctor_info = doctor_info
        self.setWindowTitle("Запланировать прием")
        self.setMinimumWidth(500)
        self.setup_ui()

    def setup_ui(self):
        """Настройка пользовательского интерфейса"""
        layout = QVBoxLayout()

        # Форма для ввода данных приема
        form_group = QGroupBox("Данные приема")
        form_layout = QFormLayout()

        # Выбор пациента
        self.patient_combo = QComboBox()
        for patient in self.patients:
            self.patient_combo.addItem(patient['full_name'], patient['id'])
        form_layout.addRow("Пациент:", self.patient_combo)

        # Дата приема
        self.appointment_date = QDateEdit()
        self.appointment_date.setDate(datetime.now().date())
        self.appointment_date.setCalendarPopup(True)
        self.appointment_date.setDisplayFormat("dd.MM.yyyy")
        form_layout.addRow("Дата:", self.appointment_date)

        # Время приема (выпадающий список с интервалами)
        self.time_combo = QComboBox()
        # Генерируем временные слоты с 9:00 до 18:00 с шагом 30 минут
        self.generate_time_slots()
        form_layout.addRow("Время:", self.time_combo)

        # Примечания
        self.notes = QTextEdit()
        self.notes.setMaximumHeight(100)
        form_layout.addRow("Примечания:", self.notes)

        form_group.setLayout(form_layout)
        layout.addWidget(form_group)

        # Кнопки действий
        buttons_layout = QHBoxLayout()

        save_button = QPushButton("Сохранить")
        save_button.setObjectName("success")
        save_button.clicked.connect(self.save_appointment)
        buttons_layout.addWidget(save_button)

        cancel_button = QPushButton("Отмена")
        cancel_button.setObjectName("secondary")
        cancel_button.clicked.connect(self.reject)
        buttons_layout.addWidget(cancel_button)

        layout.addLayout(buttons_layout)
        self.setLayout(layout)

    def generate_time_slots(self):
        """Генерация временных слотов для выбора"""
        start_time = QTime(9, 0)  # Начало рабочего дня
        end_time = QTime(18, 0)  # Конец рабочего дня
        interval = 30  # Интервал в минутах

        current_time = start_time
        while current_time <= end_time:
            self.time_combo.addItem(current_time.toString("HH:mm"), current_time)
            current_time = current_time.addSecs(interval * 60)

    def save_appointment(self):
        """Сохранение приема в базе данных"""
        patient_id = self.patient_combo.currentData()
        appointment_date = self.appointment_date.date()
        appointment_time = self.time_combo.currentData()

        # Комбинируем дату и время
        appointment_datetime = datetime(
            appointment_date.year(),
            appointment_date.month(),
            appointment_date.day(),
            appointment_time.hour(),
            appointment_time.minute()
        )
        appointment_str = appointment_datetime.strftime("%Y-%m-%d %H:%M")

        notes = self.notes.toPlainText().strip()

        if not patient_id or not appointment_date or not appointment_time:
            QMessageBox.warning(self, "Предупреждение", "Пожалуйста, заполните все обязательные поля.")
            return

        try:
            query = """
                INSERT INTO appointments (doctor_id, patient_id, appointment_date, status, notes)
                VALUES (?, ?, ?, 'scheduled', ?)
            """
            params = (self.doctor_info['id'], patient_id, appointment_str, notes)
            db.execute_query(query, params)
            QMessageBox.information(self, "Успех", "Прием успешно запланирован")
            self.accept()
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось запланировать прием: {str(e)}")


class DoctorWindow(QMainWindow):
    """Главное окно интерфейса врача"""
    logout_signal = Signal()

    def __init__(self, user_data):
        super().__init__()
        self.user_data = user_data

        # Получаем информацию о враче
        self.doctor_info = db.get_doctor_by_user_id(user_data['id'])

        self.setWindowTitle(f"Медицинский центр - ДЦ А-Линия Чертаново")
        self.setMinimumSize(800, 600)
        self.setWindowIcon(QIcon("aliniya.png"))

        # Применяем глобальный стиль
        self.setStyleSheet(GLOBAL_STYLESHEET)

        self.setup_ui()

    def setup_ui(self):
        """Настройка пользовательского интерфейса"""
        # Создание центрального виджета
        central_widget = QWidget()
        self.setCentralWidget(central_widget)

        # Основной layout
        main_layout = QVBoxLayout(central_widget)

        # Верхняя панель с информацией о пользователе и кнопками
        top_panel = QHBoxLayout()

        user_info = QLabel(f"Врач: {self.user_data['full_name']}")
        user_info.setFont(QFont("Arial", 12, QFont.Bold))
        top_panel.addWidget(user_info)

        if self.doctor_info:
            specialization = QLabel(f"Специализация: {self.doctor_info['specialization']}")
            top_panel.addWidget(specialization)

        top_panel.addStretch()

        logout_button = QPushButton("Выйти")
        logout_button.setObjectName("danger")
        logout_button.clicked.connect(self.logout)
        top_panel.addWidget(logout_button)

        main_layout.addLayout(top_panel)

        # Добавление разделительной линии
        separator = QFrame()
        separator.setFrameShape(QFrame.HLine)
        separator.setFrameShadow(QFrame.Sunken)
        main_layout.addWidget(separator)

        # Создание вкладок
        self.tab_widget = QTabWidget()

        # Вкладка расписания
        self.schedule_tab = QWidget()
        self.setup_schedule_tab()
        self.tab_widget.addTab(self.schedule_tab, "Расписание")

        # Вкладка анализов
        self.analysis_tab = QWidget()
        self.setup_analysis_tab()
        self.tab_widget.addTab(self.analysis_tab, "Анализы")

        # Вкладка рецептов
        self.prescription_tab = QWidget()
        self.setup_prescription_tab()
        self.tab_widget.addTab(self.prescription_tab, "Рецепты")

        main_layout.addWidget(self.tab_widget)

    def setup_schedule_tab(self):
        """Настройка вкладки расписания"""
        layout = QVBoxLayout(self.schedule_tab)

        # Верхняя панель с фильтрами и кнопкой планирования
        filter_layout = QHBoxLayout()

        date_label = QLabel("Дата:")
        self.date_filter = QDateEdit()
        self.date_filter.setDate(QDate.currentDate())
        self.date_filter.setCalendarPopup(True)

        apply_filter_button = QPushButton("Применить")
        apply_filter_button.setObjectName("primary")
        apply_filter_button.clicked.connect(self.load_schedule)

        schedule_button = QPushButton("Запланировать прием")
        schedule_button.setObjectName("success")
        schedule_button.clicked.connect(self.schedule_appointment)

        filter_layout.addWidget(date_label)
        filter_layout.addWidget(self.date_filter)
        filter_layout.addWidget(apply_filter_button)
        filter_layout.addWidget(schedule_button)
        filter_layout.addStretch()

        layout.addLayout(filter_layout)

        # Таблица с расписанием
        self.schedule_table = QTableWidget()
        self.schedule_table.setColumnCount(5)
        self.schedule_table.setHorizontalHeaderLabels([
            "Дата Время", "Пациент", "Статус", "Примечания", "Действия"
        ])
        self.schedule_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.schedule_table.setSortingEnabled(True)

        layout.addWidget(self.schedule_table)

        # Загрузка расписания
        self.load_schedule()

    def schedule_appointment(self):
        """Открытие диалога для планирования нового приема"""
        patients = db.get_all_patients()
        dialog = ScheduleAppointmentDialog(patients, self.doctor_info, self)
        result = dialog.exec()
        if result == QDialog.Accepted:
            self.load_schedule()

    def setup_analysis_tab(self):
        """Настройка вкладки анализов"""
        layout = QVBoxLayout(self.analysis_tab)

        # Верхняя панель с фильтрами
        filter_layout = QHBoxLayout()

        patient_label = QLabel("Пациент:")
        self.patient_filter = QComboBox()
        self.patient_filter.addItem("Все пациенты", None)

        # Загрузка списка пациентов
        patients = db.get_all_patients()
        for patient in patients:
            self.patient_filter.addItem(patient['full_name'], patient['id'])

        date_label = QLabel("Период:")
        self.start_date_filter = QDateEdit()
        self.start_date_filter.setDate(QDate.currentDate().addDays(-30))
        self.start_date_filter.setCalendarPopup(True)

        date_to_label = QLabel("по")

        self.end_date_filter = QDateEdit()
        self.end_date_filter.setDate(QDate.currentDate())
        self.end_date_filter.setCalendarPopup(True)

        apply_filter_button = QPushButton("Применить")
        apply_filter_button.setObjectName("primary")
        apply_filter_button.clicked.connect(self.load_analysis_results)

        filter_layout.addWidget(patient_label)
        filter_layout.addWidget(self.patient_filter)
        filter_layout.addWidget(date_label)
        filter_layout.addWidget(self.start_date_filter)
        filter_layout.addWidget(date_to_label)
        filter_layout.addWidget(self.end_date_filter)
        filter_layout.addWidget(apply_filter_button)
        filter_layout.addStretch()

        layout.addLayout(filter_layout)

        # Таблица с результатами анализов
        self.analysis_table = QTableWidget()
        self.analysis_table.setColumnCount(6)
        self.analysis_table.setHorizontalHeaderLabels([
            "Дата", "Пациент", "Тип анализа", "Статус", "Лаборант", "Действия"
        ])
        self.analysis_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.analysis_table.setSortingEnabled(True)

        layout.addWidget(self.analysis_table)

        # Загрузка результатов анализов
        self.load_analysis_results()

    def setup_prescription_tab(self):
        """Настройка вкладки рецептов"""
        layout = QVBoxLayout(self.prescription_tab)

        # Верхняя панель с кнопкой создания рецепта
        top_layout = QHBoxLayout()
        create_button = QPushButton("Выписать рецепт")
        create_button.setObjectName("success")
        create_button.clicked.connect(self.create_prescription)
        top_layout.addWidget(create_button)
        top_layout.addStretch()
        layout.addLayout(top_layout)

        # Таблица с рецептами
        self.prescription_table = QTableWidget()
        self.prescription_table.setColumnCount(6)
        self.prescription_table.setHorizontalHeaderLabels([
            "Дата выдачи", "Пациент", "Лекарство", "Дозировка", "Инструкции", "Действия"
        ])
        # Отключаем автоматическое растягивание столбцов
        self.prescription_table.horizontalHeader().setSectionResizeMode(QHeaderView.Interactive)
        # Разрешаем пользователю изменять размеры строк
        self.prescription_table.verticalHeader().setSectionResizeMode(QHeaderView.Interactive)
        # Включаем сортировку
        self.prescription_table.setSortingEnabled(True)
        # Разрешаем выбор только одной строки
        self.prescription_table.setSelectionMode(QTableWidget.SingleSelection)
        # Включаем отображение сетки для лучшей читаемости
        self.prescription_table.setShowGrid(True)
        layout.addWidget(self.prescription_table)

        # Загрузка рецептов
        self.load_prescriptions()

    def load_schedule(self):
        """Загрузка расписания врача"""
        if not self.doctor_info:
            QMessageBox.warning(self, "Предупреждение",
                                "Нет информации о враче. Пожалуйста, обратитесь к администратору.")
            return

        # Получение даты из фильтра
        filter_date = self.date_filter.date().toString("yyyy-MM-dd")

        # Получение расписания из базы данных
        try:
            appointments = db.fetch_all("""
                SELECT a.*, p.full_name as patient_name 
                FROM appointments a
                JOIN patients p ON a.patient_id = p.id
                WHERE a.doctor_id = ?
                ORDER BY a.appointment_date
            """, (self.doctor_info['id'],))

            # Фильтрация по дате на стороне приложения
            filtered_appointments = []
            for appointment in appointments:
                appointment_date = appointment['appointment_date'].split()[0]  # Получаем только дату без времени
                if appointment_date == filter_date:
                    filtered_appointments.append(appointment)

            # Заполнение таблицы
            self.schedule_table.setRowCount(len(filtered_appointments))

            for row, appointment in enumerate(filtered_appointments):
                # Дата и время
                appointment_datetime = appointment['appointment_date']
                try:
                    if isinstance(appointment_datetime, str):
                        date_time = datetime.strptime(appointment_datetime, "%Y-%m-%d %H:%M:%S")
                        date_item = QTableWidgetItem(date_time.strftime("%d.%m.%Y %H:%M"))
                    else:
                        date_item = QTableWidgetItem(appointment_datetime.strftime("%d.%m.%Y %H:%M"))
                except Exception as e:
                    date_item = QTableWidgetItem(str(appointment_datetime))

                self.schedule_table.setItem(row, 0, date_item)

                # Пациент
                patient_item = QTableWidgetItem(appointment['patient_name'])
                self.schedule_table.setItem(row, 1, patient_item)

                # Статус
                status_text = {
                    'scheduled': 'Запланирован',
                    'completed': 'Завершен',
                    'cancelled': 'Отменен'
                }.get(appointment['status'], appointment['status'])

                status_item = QTableWidgetItem(status_text)

                # Цветовое выделение статуса
                if appointment['status'] == 'scheduled':
                    status_item.setBackground(QColor("#ffc107"))  # Желтый
                elif appointment['status'] == 'completed':
                    status_item.setBackground(QColor("#28a745"))  # Зеленый
                elif appointment['status'] == 'cancelled':
                    status_item.setBackground(QColor("#dc3545"))  # Красный

                self.schedule_table.setItem(row, 2, status_item)

                # Примечания
                notes_item = QTableWidgetItem(appointment['notes'] if appointment['notes'] else "")
                self.schedule_table.setItem(row, 3, notes_item)

                # Кнопка действий
                view_button = QPushButton("Просмотр")
                view_button.setObjectName("primary")
                view_button.clicked.connect(lambda checked, a=appointment: self.view_appointment_details(a))
                self.schedule_table.setCellWidget(row, 4, view_button)

            # Установка высоты строк
            for row in range(self.schedule_table.rowCount()):
                self.schedule_table.setRowHeight(row, 60)  # Увеличили высоту до 60 пикселей

            self.schedule_table.resizeColumnsToContents()
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось загрузить расписание: {str(e)}")

    def load_analysis_results(self):
        """Загрузка результатов анализов"""
        # Получение параметров фильтрации
        patient_id = self.patient_filter.currentData()
        start_date = self.start_date_filter.date().toString("yyyy-MM-dd")
        end_date = self.end_date_filter.date().toString("yyyy-MM-dd")

        try:
            # Базовый запрос без фильтра по дате
            query = """
                SELECT ar.*, at.name as analysis_name, p.full_name as patient_name, u.full_name as lab_technician_name
                FROM analysis_results ar
                JOIN analysis_types at ON ar.analysis_type_id = at.id
                JOIN patients p ON ar.patient_id = p.id
                JOIN users u ON ar.lab_user_id = u.id
            """
            params = []

            # Добавление фильтра по пациенту, если выбран
            if patient_id:
                query += " WHERE ar.patient_id = ?"
                params.append(patient_id)

            query += " ORDER BY ar.result_date DESC"

            # Получение результатов анализов
            all_analysis_results = db.fetch_all(query, params)

            # Фильтрация по дате на стороне приложения
            filtered_results = []
            for result in all_analysis_results:
                try:
                    result_date_str = result['result_date']
                    if isinstance(result_date_str, str):
                        result_date = result_date_str.split()[0]  # Получаем только дату без времени
                    else:
                        result_date = result_date_str.strftime('%Y-%m-%d')

                    if start_date <= result_date <= end_date:
                        filtered_results.append(result)
                except Exception as e:
                    # Включаем результат, если не можем определить дату
                    filtered_results.append(result)

            # Заполнение таблицы
            self.analysis_table.setRowCount(len(filtered_results))

            for row, result in enumerate(filtered_results):
                # Дата (столбец 0)
                try:
                    date_str = result['result_date']
                    if isinstance(date_str, str):
                        try:
                            date_time = datetime.strptime(date_str, "%Y-%m-%d %H:%M:%S")
                            date_item = QTableWidgetItem(date_time.strftime('%d.%m.%Y %H:%M'))
                        except:
                            date_item = QTableWidgetItem(date_str)
                    else:
                        date_item = QTableWidgetItem(date_str.strftime('%d.%m.%Y %H:%M'))
                except Exception as e:
                    date_item = QTableWidgetItem(str(result.get('result_date', '')))

                self.analysis_table.setItem(row, 0, date_item)

                # Пациент (столбец 1)
                self.analysis_table.setItem(row, 1, QTableWidgetItem(result['patient_name']))

                # Тип анализа (столбец 2)
                self.analysis_table.setItem(row, 2, QTableWidgetItem(result['analysis_name']))

                # Статус (столбец 3)
                status_text = {
                    'pending': 'В обработке',
                    'completed': 'Выполнен',
                    'sent': 'Отправлен'
                }.get(result['status'], result['status'])

                status_item = QTableWidgetItem(status_text)

                # Цветовое выделение статуса
                if result['status'] == 'completed':
                    status_item.setBackground(QColor("#28a745"))  # Зеленый
                elif result['status'] == 'pending':
                    status_item.setBackground(QColor("#ffc107"))  # Желтый
                elif result['status'] == 'sent':
                    status_item.setBackground(QColor("#17a2b8"))  # Голубой

                self.analysis_table.setItem(row, 3, status_item)

                # Лаборант (столбец 4)
                self.analysis_table.setItem(row, 4, QTableWidgetItem(result['lab_technician_name']))

                # Кнопка действий (столбец 5)
                view_button = QPushButton("Просмотр")
                view_button.setObjectName("primary")
                view_button.clicked.connect(lambda checked, r=result: self.view_analysis_details(r))
                self.analysis_table.setCellWidget(row, 5, view_button)

            # Установка высоты строк
            for row in range(self.analysis_table.rowCount()):
                self.analysis_table.setRowHeight(row, 60)  # Увеличили высоту до 60 пикселей

            self.analysis_table.resizeColumnsToContents()
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось загрузить результаты анализов: {str(e)}")

    def load_prescriptions(self):
        """Загрузка выписанных рецептов"""
        try:
            prescriptions = db.fetch_all("""
                SELECT p.*, pt.full_name as patient_name
                FROM prescriptions p
                JOIN patients pt ON p.patient_id = pt.id
                WHERE p.doctor_id = ?
                ORDER BY p.issue_date DESC
            """, (self.doctor_info['id'],))

            self.prescription_table.setRowCount(len(prescriptions))

            for row, prescription in enumerate(prescriptions):
                # Дата выдачи
                date_str = prescription['issue_date']
                if isinstance(date_str, str):
                    try:
                        date_time = datetime.strptime(date_str, "%Y-%m-%d")
                        date_item = QTableWidgetItem(date_time.strftime('%d.%m.%Y'))
                    except:
                        date_item = QTableWidgetItem(date_str)
                else:
                    date_item = QTableWidgetItem(date_str.strftime('%d.%m.%Y'))
                self.prescription_table.setItem(row, 0, date_item)

                # Пациент
                self.prescription_table.setItem(row, 1, QTableWidgetItem(prescription['patient_name']))

                # Лекарства
                medications = db.get_prescription_medications(prescription['id'])
                medication_names = ", ".join([med['medication_name'] for med in medications])
                self.prescription_table.setItem(row, 2, QTableWidgetItem(medication_names))

                # Дозировки
                dosages = ", ".join([med['dosage'] for med in medications])
                self.prescription_table.setItem(row, 3, QTableWidgetItem(dosages))

                # Инструкции
                instructions = ", ".join([med['instructions'] or "" for med in medications])
                instructions_item = QTableWidgetItem(instructions)
                instructions_item.setToolTip(instructions)
                self.prescription_table.setItem(row, 4, instructions_item)

                # Кнопки действий
                actions_layout = QHBoxLayout()
                view_button = QPushButton("Word")
                view_button.setObjectName("primary")
                view_button.clicked.connect(lambda checked, pr=prescription: self.export_to_word(pr))
                actions_layout.addWidget(view_button)

                actions_widget = QWidget()
                actions_widget.setLayout(actions_layout)
                self.prescription_table.setCellWidget(row, 5, actions_widget)

                # Устанавливаем высоту строки
                self.prescription_table.setRowHeight(row, 60)

            # Удаляем вызов resizeColumnsToContents, чтобы пользователь мог сам настраивать ширину столбцов
            # self.prescription_table.resizeColumnsToContents()
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось загрузить рецепты: {str(e)}")

    def create_prescription(self):
        """Открытие диалога для создания нового рецепта"""
        patients = db.get_all_patients()
        dialog = PrescriptionDialog(patients, self.doctor_info, self)
        result = dialog.exec()
        if result == QDialog.Accepted:
            self.load_prescriptions()

    def export_to_word(self, prescription):
        """Экспорт рецепта в Word документ"""
        try:
            doc = docx.Document()
            # Настройка стилей документа
            style = doc.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(12)

            # Шапка документа
            doc.add_heading('Рецепт', level=1).alignment = 1  # Центрированное заглавие
            doc.add_paragraph('Медицинский центр "А-Линия Чертаново"', style='Normal').alignment = 1
            doc.add_paragraph('г. Москва, ул. Чертановская, д. 7, тел.: +7 (495) 123-45-67',
                              style='Normal').alignment = 1
            doc.add_paragraph().add_run().add_break()  # Пустая строка

            # Информация о враче
            doc.add_paragraph(f'Врач: {self.user_data["full_name"]}', style='Normal')
            doc.add_paragraph(f'Специализация: {self.doctor_info["specialization"]}', style='Normal')

            # Информация о пациенте
            patient = db.get_patient(prescription['patient_id'])
            doc.add_paragraph(f'Пациент: {patient["full_name"]}', style='Normal')
            doc.add_paragraph(f'Дата рождения: {patient["birth_date"]}', style='Normal')
            doc.add_paragraph(f'Дата выдачи: {prescription["issue_date"]}', style='Normal')
            doc.add_paragraph()

            # Таблица с лекарствами
            medications = db.get_prescription_medications(prescription['id'])
            if medications:
                table = doc.add_table(rows=len(medications) + 1, cols=3)
                table.style = 'Table Grid'
                # Заголовки таблицы
                header_cells = table.rows[0].cells
                header_cells[0].text = 'Лекарство'
                header_cells[1].text = 'Дозировка'
                header_cells[2].text = 'Инструкции по применению'
                for cell in header_cells:
                    cell.paragraphs[0].style.font.bold = True
                    cell.paragraphs[0].alignment = 1

                # Заполнение таблицы
                for i, med in enumerate(medications, 1):
                    row_cells = table.rows[i].cells
                    row_cells[0].text = med['medication_name']
                    row_cells[1].text = med['dosage']
                    row_cells[2].text = med['instructions'] or 'Нет инструкций'

            else:
                doc.add_paragraph('Лекарства не назначены', style='Normal')

            # Подпись
            doc.add_paragraph()
            doc.add_paragraph('____________________', style='Normal').alignment = 2  # Справа
            doc.add_paragraph(f'Подпись врача: {self.user_data["full_name"]}', style='Normal').alignment = 2
            doc.add_paragraph(f'Дата: {prescription["issue_date"]}', style='Normal').alignment = 2

            # Сохранение документа
            filename = f"Рецепт_{patient['full_name']}_{prescription['issue_date']}.docx"
            filename = filename.replace(" ", "_").replace(":", "-")
            doc.save(filename)
            QMessageBox.information(self, "Успех", f"Рецепт сохранен как {filename}")
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось экспортировать рецепт в Word: {str(e)}")

    def view_appointment_details(self, appointment):
        """Просмотр деталей приема"""
        dialog = AppointmentDetailsDialog(appointment, self)
        result = dialog.exec()

        # Обновление таблицы после изменения статуса
        if result == QDialog.Accepted:
            self.load_schedule()

    def view_analysis_details(self, analysis):
        """Просмотр деталей анализа"""
        dialog = AnalysisDetailsDialog(analysis, self)
        dialog.exec()

    def logout(self):
        """Выход из системы"""
        reply = QMessageBox.question(
            self,
            "Подтверждение выхода",
            "Вы уверены, что хотите выйти?",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No
        )

        if reply == QMessageBox.Yes:
            self.logout_signal.emit()


if __name__ == "__main__":
    from PySide6.QtWidgets import QApplication

    app = QApplication(sys.argv)
    app.setStyleSheet(GLOBAL_STYLESHEET)  # Применяем стиль ко всему приложению

    # Тестовое подключение к базе данных
    db.connect("1")  # Пароль для базы данных

    # Тестовый пользователь
    test_user = {
        'id': 2,
        'username': 'doctor1',
        'full_name': 'Петров Иван Сергеевич',
        'role': 'doctor'
    }

    window = DoctorWindow(test_user)
    window.show()

    sys.exit(app.exec())