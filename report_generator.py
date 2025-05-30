import os
import json
from datetime import datetime
import docx
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PySide6.QtWidgets import QFileDialog, QMessageBox
from PySide6.QtCore import QDate

from database_connection import DatabaseConnection

# Создаем экземпляр менеджера БД
db = DatabaseConnection()
db.connect()

def export_analysis_to_word(result_id, parent_widget):
    """Экспорт результата анализа в Word"""
    # Получаем данные о результате анализа
    query = """
        SELECT ar.id, ar.result_date as date, p.full_name as patient_name, p.birth_date, 
               p.gender, p.phone, at.name as analysis_type, ar.result_data, 
               ar.status, u.full_name as lab_technician
        FROM analysis_results ar
        JOIN patients p ON ar.patient_id = p.id
        JOIN analysis_types at ON ar.analysis_type_id = at.id
        JOIN users u ON ar.lab_user_id = u.id
        WHERE ar.id = ?
    """
    result = db.fetch_one(query, (result_id,))
    
    if not result:
        QMessageBox.warning(parent_widget, "Ошибка", "Результат анализа не найден.")
        return
    
    # Создаем документ Word
    doc = docx.Document()
    
    # Настройка полей документа
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)
    
    # Заголовок документа
    header = doc.add_heading('РЕЗУЛЬТАТ АНАЛИЗА', level=1)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Информация о пациенте
    doc.add_heading('Информация о пациенте', level=2)
    
    patient_table = doc.add_table(rows=5, cols=2)
    patient_table.style = 'Table Grid'
    
    # Данные пациента
    cells = [
        ('ФИО пациента:', result['patient_name']),
        ('Дата рождения:', result['birth_date']),
        ('Пол:', result['gender']),
        ('Телефон:', result['phone']),
        ('Дата анализа:', result['date'])
    ]
    
    for i, (label, value) in enumerate(cells):
        patient_table.cell(i, 0).text = label
        patient_table.cell(i, 1).text = str(value)
    
    # Информация об анализе
    doc.add_heading(f'Анализ: {result["analysis_type"]}', level=2)
    
    # Данные результатов анализа
    result_data = json.loads(result['result_data']) if result['result_data'] else {}
    
    if result_data:
        results_table = doc.add_table(rows=len(result_data) + 1, cols=3)
        results_table.style = 'Table Grid'
        
        # Заголовки таблицы
        headers = ['Параметр', 'Значение', 'Нормальные значения']
        for i, header in enumerate(headers):
            results_table.cell(0, i).text = header
        
        # Заполняем таблицу результатами
        for i, (param, value) in enumerate(result_data.items(), 1):
            results_table.cell(i, 0).text = param
            results_table.cell(i, 1).text = str(value)
            
            # Получаем нормальные значения для параметра
            normal_values = get_normal_values(param)
            results_table.cell(i, 2).text = normal_values
    else:
        doc.add_paragraph('Нет данных о результатах анализа.')
    
    # Подпись лаборанта
    doc.add_paragraph('')
    doc.add_paragraph(f'Лаборант: {result["lab_technician"]}')
    doc.add_paragraph(f'Дата: {datetime.now().strftime("%d.%m.%Y")}')
    
    # Выбор места сохранения файла
    default_filename = f'Анализ_{result["analysis_type"]}_{result["patient_name"]}_{result["date"]}.docx'
    default_filename = default_filename.replace('/', '-').replace('\\', '-').replace(':', '-')
    
    file_path, _ = QFileDialog.getSaveFileName(
        parent_widget,
        "Сохранить отчет",
        default_filename,
        "Word Documents (*.docx)"
    )
    
    if file_path:
        if not file_path.endswith('.docx'):
            file_path += '.docx'
        
        try:
            doc.save(file_path)
            QMessageBox.information(parent_widget, "Успех", "Отчет успешно сохранен.")
            return file_path
        except Exception as e:
            QMessageBox.critical(parent_widget, "Ошибка", f"Не удалось сохранить файл: {str(e)}")
    
    return None

def export_all_analyses_to_word(parent_widget, filters=None):
    """Экспорт всех результатов анализов в Word с учетом фильтров"""
    # Базовый запрос
    query = """
        SELECT ar.id, ar.result_date as date, p.full_name as patient_name, p.birth_date, 
               p.gender, p.phone, at.name as analysis_type, ar.result_data, 
               ar.status, u.full_name as lab_technician
        FROM analysis_results ar
        JOIN patients p ON ar.patient_id = p.id
        JOIN analysis_types at ON ar.analysis_type_id = at.id
        JOIN users u ON ar.lab_user_id = u.id
        WHERE 1=1
    """
    
    params = []
    
    # Применяем фильтры, если они указаны
    if filters:
        if filters.get('patient_id'):
            query += " AND p.id = ?"
            params.append(filters['patient_id'])
        
        if filters.get('analysis_type_id'):
            query += " AND at.id = ?"
            params.append(filters['analysis_type_id'])
        
        if filters.get('from_date'):
            query += " AND ar.result_date >= ?"
            params.append(filters['from_date'])
        
        if filters.get('to_date'):
            query += " AND ar.result_date <= ?"
            params.append(filters['to_date'])
        
        if filters.get('status'):
            query += " AND ar.status = ?"
            params.append(filters['status'])
    
    # Получаем результаты
    results = db.fetch_all(query, tuple(params))
    
    if not results:
        QMessageBox.warning(parent_widget, "Внимание", "Нет результатов анализов для экспорта.")
        return
    
    # Создаем документ Word
    doc = docx.Document()
    
    # Настройка полей документа
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)
    
    # Заголовок документа
    header = doc.add_heading('ОТЧЕТ ПО РЕЗУЛЬТАТАМ АНАЛИЗОВ', level=1)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Добавляем дату отчета
    date_paragraph = doc.add_paragraph(f'Дата формирования: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Общая информация
    doc.add_paragraph(f'Количество результатов: {len(results)}')
    
    # Проходим по каждому результату
    for i, result in enumerate(results, 1):
        doc.add_heading(f'{i}. {result["analysis_type"]} - {result["patient_name"]}', level=2)
        
        # Информация о пациенте
        patient_info = doc.add_paragraph()
        patient_info.add_run('Пациент: ').bold = True
        patient_info.add_run(f'{result["patient_name"]}, {result["birth_date"]}, {result["gender"]}, тел: {result["phone"]}')
        
        date_info = doc.add_paragraph()
        date_info.add_run('Дата анализа: ').bold = True
        date_info.add_run(result['date'])
        
        status_info = doc.add_paragraph()
        status_info.add_run('Статус: ').bold = True
        status_info.add_run(result['status'])
        
        lab_tech_info = doc.add_paragraph()
        lab_tech_info.add_run('Лаборант: ').bold = True
        lab_tech_info.add_run(result['lab_technician'])
        
        # Данные результатов анализа
        result_data = json.loads(result['result_data']) if result['result_data'] else {}
        
        if result_data:
            doc.add_heading('Результаты анализа:', level=3)
            results_table = doc.add_table(rows=len(result_data) + 1, cols=3)
            results_table.style = 'Table Grid'
            
            # Заголовки таблицы
            headers = ['Параметр', 'Значение', 'Нормальные значения']
            for j, header in enumerate(headers):
                results_table.cell(0, j).text = header
            
            # Заполняем таблицу результатами
            for j, (param, value) in enumerate(result_data.items(), 1):
                results_table.cell(j, 0).text = param
                results_table.cell(j, 1).text = str(value)
                
                # Получаем нормальные значения для параметра
                normal_values = get_normal_values(param)
                results_table.cell(j, 2).text = normal_values
        else:
            doc.add_paragraph('Нет данных о результатах анализа.')
        
        # Добавляем разделитель между результатами
        if i < len(results):
            doc.add_paragraph('_' * 50)
    
    # Выбор места сохранения файла
    date_str = datetime.now().strftime("%Y%m%d_%H%M")
    default_filename = f'Отчет_по_анализам_{date_str}.docx'
    
    file_path, _ = QFileDialog.getSaveFileName(
        parent_widget,
        "Сохранить отчет",
        default_filename,
        "Word Documents (*.docx)"
    )
    
    if file_path:
        if not file_path.endswith('.docx'):
            file_path += '.docx'
        
        try:
            doc.save(file_path)
            QMessageBox.information(parent_widget, "Успех", "Отчет успешно сохранен.")
            return file_path
        except Exception as e:
            QMessageBox.critical(parent_widget, "Ошибка", f"Не удалось сохранить файл: {str(e)}")
    
    return None

def get_normal_values(parameter):
    """Получение нормальных значений для параметра анализа"""
    # В реальном приложении это должно быть получено из БД
    # Здесь приведены условные примеры
    normal_values = {
        'Гемоглобин': '120-160 г/л (Ж), 130-170 г/л (М)',
        'Эритроциты': '3,8-5,2 ×10^12/л (Ж), 4,2-5,6 ×10^12/л (М)',
        'Лейкоциты': '4,0-9,0 ×10^9/л',
        'Глюкоза': '3,3-5,5 ммоль/л',
        'Холестерин': '3,6-5,2 ммоль/л',
        'Триглицериды': '0,45-1,7 ммоль/л',
        'АЛТ': '0-33 Ед/л (Ж), 0-41 Ед/л (М)',
        'АСТ': '0-32 Ед/л (Ж), 0-40 Ед/л (М)',
        'Билирубин общий': '3,4-20,5 мкмоль/л',
        'Креатинин': '44-80 мкмоль/л (Ж), 62-106 мкмоль/л (М)',
        'Мочевина': '2,5-6,7 ммоль/л',
        'Мочевая кислота': '154-357 мкмоль/л (Ж), 202-416 мкмоль/л (М)'
    }
    
    return normal_values.get(parameter, 'Нет данных') 